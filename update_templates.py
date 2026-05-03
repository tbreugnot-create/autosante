"""
Script de mise à jour des templates de bons de consommation santé.
Ajoute les placeholders <<DETTE_TOTALE>> et <<RETENUE_MOIS>> après <<NOM>>.

Usage :
    cd "/Users/thomas/Documents/Claude/Projects/Pôle Santé"
    python update_templates.py

Les fichiers modifiés sont sauvegardés avec le suffixe _updated.docx
(les originaux ne sont pas écrasés).
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import lxml.etree as etree

TEMPLATES = [
    "Template - Bon de consommation santé - Employeur",
    "Template - Bon de consommation santé - Employé",
]

# Chercher les fichiers dans le dossier courant par mot-clé
def find_template(keyword):
    """Cherche un .docx dont le nom contient 'keyword', en ignorant la casse et les accents."""
    import unicodedata

    def normalize(s):
        return unicodedata.normalize("NFC", s).lower()

    kw = normalize(keyword)
    candidates = []
    for fname in os.listdir("."):
        if not fname.lower().endswith(".docx"):
            continue
        if "_updated" in fname:
            continue
        if kw in normalize(fname):
            candidates.append(fname)

    if not candidates:
        return None
    # Prendre le plus récent si plusieurs
    candidates.sort(key=lambda f: os.path.getmtime(f), reverse=True)
    return candidates[0]


def add_debt_section(doc_path, out_path):
    doc = Document(doc_path)

    # Trouver le paragraphe <<NOM>>
    nom_para_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "<<NOM>>" in para.text:
            nom_para_idx = i
            break

    if nom_para_idx is None:
        print(f"  ⚠️  Placeholder <<NOM>> non trouvé dans {doc_path}. Skipped.")
        return

    # Construire les nouveaux paragraphes à insérer après <<NOM>>
    # On les insère dans l'XML directement (python-docx ne supporte pas insert_after)
    nom_para = doc.paragraphs[nom_para_idx]
    parent = nom_para._element.getparent()
    nom_idx_in_parent = list(parent).index(nom_para._element)

    def make_info_para(label: str, placeholder: str, bold_label: bool = True) -> etree._Element:
        """Crée un paragraphe XML de la forme 'label : <<PLACEHOLDER>>'."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)
        p.append(pPr)

        # Run label (gras)
        r_label = OxmlElement("w:r")
        rPr_label = OxmlElement("w:rPr")
        if bold_label:
            b = OxmlElement("w:b")
            rPr_label.append(b)
        # Taille 10pt
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "20")
        rPr_label.append(sz)
        rPr_label.append(szCs)
        r_label.append(rPr_label)
        t_label = OxmlElement("w:t")
        t_label.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_label.text = label + " : "
        r_label.append(t_label)
        p.append(r_label)

        # Run placeholder
        r_ph = OxmlElement("w:r")
        rPr_ph = OxmlElement("w:rPr")
        sz2 = OxmlElement("w:sz")
        sz2.set(qn("w:val"), "20")
        szCs2 = OxmlElement("w:szCs")
        szCs2.set(qn("w:val"), "20")
        rPr_ph.append(sz2)
        rPr_ph.append(szCs2)
        r_ph.append(rPr_ph)
        t_ph = OxmlElement("w:t")
        t_ph.text = placeholder
        r_ph.append(t_ph)
        p.append(r_ph)

        return p

    # Vérifier si les placeholders existent déjà
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "<<DETTE_TOTALE>>" in full_text and "<<RETENUE_MOIS>>" in full_text:
        print(f"  ℹ️  Placeholders déjà présents dans {doc_path}. Skipped.")
        return

    # Insérer un paragraphe vide + dette + retenue après <<NOM>>
    insert_after = nom_idx_in_parent

    # Paragraphe vide séparateur
    sep = etree.SubElement(parent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
    parent.remove(sep)
    parent.insert(insert_after + 1, sep)

    # Paragraphe dette totale
    p_dette = make_info_para("Encours total dû (dette santé)", "<<DETTE_TOTALE>>")
    parent.insert(insert_after + 2, p_dette)

    # Paragraphe retenue
    p_retenue = make_info_para("Retenue sur salaire du mois", "<<RETENUE_MOIS>>")
    parent.insert(insert_after + 3, p_retenue)

    doc.save(out_path)
    print(f"  ✅  Sauvegardé : {out_path}")


if __name__ == "__main__":
    print("=== Mise à jour des templates de bons de consommation ===\n")
    for tpl_prefix in TEMPLATES:
        is_employeur = "Employeur" in tpl_prefix

        # Chercher tous les .docx du dossier
        import unicodedata
        def _norm(s): return unicodedata.normalize("NFC", s).lower()

        candidates = []
        for fname in os.listdir("."):
            if not fname.lower().endswith(".docx"):
                continue
            if "_updated" in fname:
                continue
            fn = _norm(fname)
            if is_employeur:
                if "employeur" in fn:
                    candidates.append(fname)
            else:
                # Employé uniquement — exclure Employeur
                if "employ" in fn and "employeur" not in fn:
                    candidates.append(fname)

        if not candidates:
            label = "Employeur" if is_employeur else "Employé"
            print(f"❌ Fichier non trouvé pour : {label}")
            print("   Télécharge les templates depuis Google Drive dans ce dossier et relance le script.")
            continue

        candidates.sort(key=lambda f: os.path.getmtime(f), reverse=True)
        path = candidates[0]
        label = "Employeur" if is_employeur else "Employe"
        out = f"./Template - Bon de consommation santé - {label}_updated.docx"
        print(f"Traitement : {path}")
        add_debt_section(path, out)

    print("\nTerminé. Les fichiers _updated.docx sont prêts à être uploadés dans Google Drive.")
    print("Pense à uploader les nouveaux templates dans l'app (section 'Templates bons').")
