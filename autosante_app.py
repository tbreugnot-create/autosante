"""
AutoSanté v2.2 — Odoo API direct · Paramètres via Google Sheets · YTD + Provision Santé
Société : DI-Africa (Congo) SA uniquement (company_id = 3)

Installation :
    pip install streamlit openpyxl python-docx

Lancement :
    streamlit run autosante_app.py
"""

import re
import io
import csv
import zipfile
import calendar
import urllib.request
import urllib.parse
from datetime import datetime
from collections import defaultdict

import streamlit as st
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import xmlrpc.client
from docx import Document

# ── CONFIGURATION ─────────────────────────────────────────────────────────
ODOO_URL   = "https://www.di-africa.com"
ODOO_DB    = "odoo-ps-psbe-di-africa-co1-main-30213273"
COMPANY_ID = 3   # DI-Africa (Congo) SA

# Identifiants Odoo — lus depuis Streamlit Secrets en prod,
# ou depuis les variables ci-dessous en local (fichier .streamlit/secrets.toml)
def _cfg(key: str, default: str) -> str:
    try:
        return st.secrets[key]
    except Exception:
        return default

ODOO_EMAIL = _cfg("ODOO_EMAIL", "")
ODOO_KEY   = _cfg("ODOO_KEY",   "")

# Google Sheet partagé pour Taux Clients + Prestataires
# → Partager la feuille "Paramètres" en "Tout le monde peut consulter"
# → Copier l'ID depuis l'URL : docs.google.com/spreadsheets/d/{ID}/edit
PARAMS_SHEET_ID = "1nkvAb3rjrqXF95K4ZYC5ZfIXqMusSesCGR8kQTw9dsw"  # Paramètres TBR

# Regex extraction client depuis le département Odoo
# "BU Congo / Clients Congo / ERoCo / Mengo"  →  "ERoCo"
_CLIENT_RE = re.compile(
    r'Clients Congo\s*/\s*([\w\d\s\-\.&()]+?)(?:\s*/\s*[\w\d\s]+\s*$|$)',
    re.IGNORECASE
)

# ── CLASSIFICATION DES ARTICLES DE FACTURATION ───────────────────────────
# Détermine la règle de partage employeur/employé selon le code article Odoo.
#
#  "salarie"         → table "Taux Clients"         (taux contractuels par client)
#  "ayants_droits"   → table "Taux Ayants Droits"   (taux spécifiques AD par client)
#  "ayants_droits_na"→ 100 % employé                (non autorisé)
#  "accident"        → 100 % employeur              (accident du travail)
#  "insurance"       → 100 % employeur              (assurance / insurance)
#  "acpe_vm"         → 100 % employeur              (ACPE / VM)
#
# Priorité de matching : du plus spécifique au plus générique.

def _classify_article(product_name: str) -> str:
    """Retourne le type d'article médical à partir du nom de produit Odoo."""
    p = (product_name or "").lower()
    if "accident" in p:
        return "accident"
    if "insurance" in p or "assurance" in p:
        return "insurance"
    if "acpe" in p or "( vm)" in p or "(vm)" in p or "acpe vm" in p:
        return "acpe_vm"
    if "ayant" in p:
        # "NA" = non autorisé → 100% employé
        if " na" in p or "/na" in p or p.endswith("na"):
            return "ayants_droits_na"
        return "ayants_droits"
    if "salarié" in p or "salarie" in p or "salari" in p:
        return "salarie"
    # Fallback : traitement standard salarié
    return "salarie"

# ── CONNEXION ODOO ────────────────────────────────────────────────────────
@st.cache_resource(show_spinner="Connexion à Odoo…")
def odoo_connect():
    common = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/common")
    uid    = common.authenticate(ODOO_DB, ODOO_EMAIL, ODOO_KEY, {})
    if not uid:
        st.error("❌ Échec de connexion à Odoo. Vérifiez les identifiants.")
        st.stop()
    models = xmlrpc.client.ServerProxy(f"{ODOO_URL}/xmlrpc/2/object")
    return uid, models


def odoo_read(models, uid, model, domain, fields, limit=10000, order=None):
    kwargs = {"fields": fields, "limit": limit}
    if order:
        kwargs["order"] = order
    return models.execute_kw(ODOO_DB, uid, ODOO_KEY, model,
                             "search_read", [domain], kwargs)


# ── EXTRACTION ODOO ───────────────────────────────────────────────────────
@st.cache_data(ttl=300, show_spinner="Récupération des factures…")
def fetch_invoice_lines(_uid, _models, date_from: str, date_to: str):
    """Lignes de facture santé DI-Africa Congo sur la période."""
    # Filtre identique au filtre Odoo manuel d'Aurice :
    # "Invoice lines contains Medical" → product_id.name ilike "Medical"
    # + Optique pour les soins optiques
    domain = [
        ["company_id",            "=",  COMPANY_ID],
        ["move_id.state",         "=",  "posted"],
        ["date",                  ">=", date_from],
        ["date",                  "<=", date_to],
        ["x_studio_employee_inv", "!=", False],
        "|",
        ["product_id.name", "ilike", "Medical"],
        ["product_id.name", "ilike", "Optique"],
    ]
    fields = [
        "move_id", "product_id", "balance",
        "partner_id", "date", "x_studio_employee_inv",
    ]
    return odoo_read(_models, _uid, "account.move.line", domain, fields)


@st.cache_data(ttl=300, show_spinner="Récupération des données YTD (Jan → mois sélectionné)…")
def fetch_invoice_lines_ytd(_uid, _models, year: int, month: int):
    """
    Lignes de facture santé DI-Africa Congo du 1er janvier au dernier jour
    du mois sélectionné — utilisé pour les cumuls YTD et le suivi provision.
    """
    date_from = f"{year}-01-01"
    last_day  = calendar.monthrange(year, month)[1]
    date_to   = f"{year}-{month:02d}-{last_day:02d}"
    domain = [
        ["company_id",            "=",  COMPANY_ID],
        ["move_id.state",         "=",  "posted"],
        ["date",                  ">=", date_from],
        ["date",                  "<=", date_to],
        ["x_studio_employee_inv", "!=", False],
        "|",
        ["product_id.name", "ilike", "Medical"],
        ["product_id.name", "ilike", "Optique"],
    ]
    fields = [
        "move_id", "product_id", "balance",
        "partner_id", "date", "x_studio_employee_inv",
    ]
    return odoo_read(_models, _uid, "account.move.line", domain, fields)


@st.cache_data(ttl=300, show_spinner="Récupération des retenues Odoo (hr.salary.attachment)…")
def fetch_salary_attachments(_uid, _models):
    """
    Lit tous les hr.salary.attachment de type 'Retraits Santé' actifs
    pour la société DI-Africa Congo.

    Retourne : {
      emp_name: {
        "odoo_id":        int,
        "description":    str,
        "employee_id":    int,
        "employee_name":  str,
        "monthly_amount": float,   # retenue mensuelle actuellement en place
        "total_amount":   float,   # dette totale à recouvrer (s'incrémente)
        "paid_amount":    float,   # déjà prélevé via paie
        "remaining":      float,   # = total_amount - paid_amount
        "date_start":     str,
        "state":          str,     # "open" | "running" | "done"
        "input_type":     str,
      }
    }
    """
    domain = [
        ["company_id", "=",    COMPANY_ID],
        ["state",      "in",   ["open", "running"]],
    ]
    fields = [
        "id", "description", "employee_id",
        "monthly_amount", "total_amount", "paid_amount",
        "date_start", "state", "other_input_type_id",
    ]
    try:
        records = odoo_read(_models, _uid, "hr.salary.attachment",
                            domain, fields, limit=5000)
    except Exception as e:
        st.warning(f"⚠️ Impossible de lire hr.salary.attachment : {e}")
        return {}

    result = {}
    for r in records:
        input_type = r.get("other_input_type_id")
        type_name  = input_type[1] if input_type else ""
        # Garder uniquement les Retraits Santé
        if not ("sant" in type_name.lower() or "retrait" in type_name.lower()):
            continue
        emp      = r.get("employee_id")
        emp_id   = emp[0] if emp else None
        emp_name = emp[1] if emp else ""
        total    = float(r.get("total_amount", 0) or 0)
        paid     = float(r.get("paid_amount",  0) or 0)
        result[emp_name] = {
            "odoo_id":        r["id"],
            "description":    r.get("description", "Santé Déduction"),
            "employee_id":    emp_id,
            "employee_name":  emp_name,
            "monthly_amount": float(r.get("monthly_amount", 0) or 0),
            "total_amount":   total,
            "paid_amount":    paid,
            "remaining":      max(0.0, total - paid),
            "date_start":     r.get("date_start", ""),
            "state":          r.get("state", ""),
            "input_type":     type_name,
        }
    return result


@st.cache_data(ttl=3600, show_spinner="Chargement des employés (actifs + archivés)…")
def fetch_employees(_uid, _models):
    """Employés DI-Africa Congo — actifs ET archivés pour garantir
    que les factures d'ex-employés sont bien rattachées."""
    rows = odoo_read(_models, _uid, "hr.employee",
                     [["company_id", "=", COMPANY_ID],
                      ["active", "in", [True, False]]],   # inclut les archivés
                     ["id", "name", "department_id", "active"], limit=5000)
    result = {}
    for r in rows:
        dept  = r["department_id"][1] if r["department_id"] else ""
        m     = _CLIENT_RE.search(dept)
        if m:
            client = m.group(1).strip()
            # Supprimer éventuel sous-site résiduel
            client = re.sub(r'\s*/.*$', '', client).strip()
        else:
            # Pas de "Clients Congo" dans le département → employé corporate DI-Africa
            client = "DI-Africa"
        result[r["id"]] = {
            "name":   r["name"],
            "dept":   dept,
            "client": client,
            "active": r.get("active", True),   # False = employé archivé
        }
    return result


# ── LECTURE PARAMÈTRES DEPUIS GOOGLE SHEETS ───────────────────────────────
@st.cache_data(ttl=600, show_spinner="Chargement des paramètres depuis Google Sheets…")
def load_params_from_gsheet(sheet_id: str) -> dict:
    """
    Lit les feuilles 'Taux Clients' et 'Prestataires' depuis le Google Sheet partagé
    via export CSV (pas d'authentification requise — partage "tout le monde peut consulter").

    Retourne :
      params["rates"]        → {client: {cc, consult_soc, consult_emp,
                                          pharma_soc, pharma_emp, optique_soc, optique_emp}}
      params["prestataires"] → {nom_upper: type}  (type = Consultation | Pharmacie | Optique)
    """
    base_url = (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}"
        "/gviz/tq?tqx=out:csv&sheet="
    )

    def fetch_sheet_csv(sheet_name: str) -> list:
        url = base_url + urllib.parse.quote(sheet_name)
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "AutoSante/2.1"})
            with urllib.request.urlopen(req, timeout=15) as resp:
                raw = resp.read().decode("utf-8")
            return list(csv.reader(raw.splitlines()))
        except Exception as e:
            st.error(f"❌ Impossible de lire la feuille « {sheet_name} » : {e}")
            st.stop()

    def _f(v, default=0.0):
        """
        Convertit une valeur de TAUX en float décimal (0.0–1.0).
        Gère : 0.8 · "80%" · "80" · "0,8" · cellule vide
        Utilisé pour consultations et pharmacies uniquement.
        """
        s = str(v).replace(",", ".").strip()
        has_pct = s.endswith("%")
        s = s.rstrip("%").strip()
        try:
            result = float(s) if s else default
            if has_pct or result > 1.0:
                result /= 100.0
            return result
        except (ValueError, TypeError):
            return default

    def _f_optique(v):
        """
        Parse la colonne Optique qui peut contenir :
        - Un PLAFOND en FCFA (ex: 110000) → retourner {"type": "cap", "val": 110000}
        - Un TAUX en % (ex: "50%", "100%") → retourner {"type": "rate", "val": 0.5}
        - Vide → None
        """
        s = str(v).replace(",", ".").strip() if v else ""
        if not s:
            return None
        has_pct = s.endswith("%")
        s_num = s.rstrip("%").strip()
        try:
            num = float(s_num)
            if has_pct:
                return {"type": "rate", "val": num / 100.0}
            elif num <= 1.0:
                return {"type": "rate", "val": num}
            else:
                # Grande valeur sans % → plafond en FCFA (ex: 110000)
                return {"type": "cap", "val": num}
        except (ValueError, TypeError):
            return None

    def _f_plafond(v):
        """
        Parse un plafond annuel de provision santé en FCFA.
        Accepte : 5000000 · "5 000 000" · "5,000,000" · vide → None
        """
        s = str(v).replace(",", "").replace(" ", "").replace(" ", "").strip()
        if not s:
            return None
        try:
            return float(s)
        except (ValueError, TypeError):
            return None

    # Taux Clients
    # Colonnes attendues dans "Taux Clients" :
    #  0=CC  1=Client  2=Consult_Soc  3=Consult_Emp  4=Pharma_Soc  5=Pharma_Emp
    #  6=Optique_Soc  7=Optique_Emp  8=Modèle  9=Plafond_Contrat  10=Plafond_Employé
    #  11=Règle_Retenue  (optionnel : "15%"|"total"|"1/3"|"1/4"|"1/5"|"fixe:25000")
    #
    # Optique_Soc peut être :
    #   - Un FCFA (ex: 110000) → plafond annuel par employé (cumulatif YTD)
    #   - Un % (ex: "80%")     → taux par facture, reliquat à charge employé
    rates = {}
    taux_rows = fetch_sheet_csv("Taux Clients")
    for row in taux_rows[1:]:   # skip header
        if len(row) < 2 or not row[1].strip():
            continue
        client = row[1].strip()
        raw_modele = row[8].strip().lower() if len(row) > 8 and row[8].strip() else "open bar"
        modele = "provision" if "provision" in raw_modele else "open bar"
        raw_regle = row[11].strip().lower() if len(row) > 11 and row[11].strip() else "15%"
        rates[client] = {
            "cc":             row[0].strip() if row[0] else "",
            "consult_soc":    _f(row[2] if len(row) > 2 else 0),
            "consult_emp":    _f(row[3] if len(row) > 3 else 0),
            "pharma_soc":     _f(row[4] if len(row) > 4 else 0),
            "pharma_emp":     _f(row[5] if len(row) > 5 else 0),
            "optique_soc":    _f_optique(row[6] if len(row) > 6 else ""),
            "optique_emp":    _f_optique(row[7] if len(row) > 7 else ""),
            "modele":         modele,                                        # "open bar" | "provision"
            "plafond":        _f_plafond(row[9]  if len(row) > 9  else ""), # Plafond contrat FCFA
            "plafond_emp":    _f_plafond(row[10] if len(row) > 10 else ""), # Plafond par employé FCFA/an
            "regle_retenue":  raw_regle,                                     # règle retenue sur salaire
        }

    # Taux Ayants Droits
    # Même structure que "Taux Clients" — colonnes identiques.
    # Si la feuille n'existe pas encore, on retourne un dict vide (pas de blocage).
    rates_ad = {}
    try:
        taux_ad_rows = fetch_sheet_csv("Taux Ayants Droits")
        for row in taux_ad_rows[1:]:
            if len(row) < 2 or not row[1].strip():
                continue
            client_ad = row[1].strip()
            raw_modele_ad = row[8].strip().lower() if len(row) > 8 and row[8].strip() else "open bar"
            modele_ad = "provision" if "provision" in raw_modele_ad else "open bar"
            raw_regle_ad = row[11].strip().lower() if len(row) > 11 and row[11].strip() else "15%"
            rates_ad[client_ad] = {
                "cc":             row[0].strip() if row[0] else "",
                "consult_soc":    _f(row[2] if len(row) > 2 else 0),
                "consult_emp":    _f(row[3] if len(row) > 3 else 0),
                "pharma_soc":     _f(row[4] if len(row) > 4 else 0),
                "pharma_emp":     _f(row[5] if len(row) > 5 else 0),
                "optique_soc":    _f_optique(row[6] if len(row) > 6 else ""),
                "optique_emp":    _f_optique(row[7] if len(row) > 7 else ""),
                "modele":         modele_ad,
                "plafond":        _f_plafond(row[9]  if len(row) > 9  else ""),
                "plafond_emp":    _f_plafond(row[10] if len(row) > 10 else ""),
                "regle_retenue":  raw_regle_ad,
            }
    except Exception:
        # Feuille absente ou erreur → pas de blocage, avertissement en UI
        rates_ad = {}

    # Prestataires
    prest = {}
    prest_rows = fetch_sheet_csv("Prestataires")
    for row in prest_rows[1:]:  # skip header
        if len(row) >= 2 and row[0].strip():
            prest[row[0].strip().upper()] = row[1].strip()

    return {"rates": rates, "rates_ad": rates_ad, "prestataires": prest}


# ── TRAITEMENT PRINCIPAL ──────────────────────────────────────────────────

def _find_rate(client: str, rates: dict) -> dict | None:
    """Recherche du taux client — exact puis partiel, insensible à la casse."""
    for k in rates:
        if k.lower() == client.lower():
            return rates[k]
    for k in rates:
        if client.lower() in k.lower() or k.lower() in client.lower():
            return rates[k]
    return None


def compute_ytd_optique_consumed(lines_prior: list, employees: dict,
                                  params: dict) -> dict:
    """
    Calcule, pour chaque employé, la part déjà payée par la société
    en optique sur les mois précédents (Jan → M-1).

    Utilisé pour appliquer correctement le plafond optique annuel par employé :
    si un employé a déjà consommé 80 000 FCFA sur son cap annuel de 110 000,
    il ne reste que 30 000 FCFA de prise en charge société pour le mois courant.

    Retourne : {employee_name: montant_soc_optique_ytd_precedents}
    """
    prests = params.get("prestataires", {})
    rates  = params.get("rates", {})
    consumed: dict[str, float] = defaultdict(float)

    for ln in lines_prior:
        emp_name     = ln["x_studio_employee_inv"][1]
        emp_id       = ln["x_studio_employee_inv"][0]
        partner_name = ln["partner_id"][1] if ln["partner_id"] else ""
        product_name = ln["product_id"][1] if ln["product_id"] else ""
        balance      = abs(float(ln["balance"] or 0))

        # Détermination type prestataire
        prest_type = prests.get(partner_name.upper())
        if prest_type is None:
            prest_type = "Optique" if "optique" in product_name.lower() else "Pharmacie"

        if prest_type != "Optique":
            continue  # On ne s'intéresse qu'à l'optique

        emp_info = employees.get(emp_id, {"client": "DI-Africa"})
        client   = emp_info.get("client", "")
        rate     = _find_rate(client, rates)
        if rate is None:
            continue

        opt_soc = rate.get("optique_soc")
        if opt_soc is None:
            continue

        # Calcul de la part société payée sur les mois précédents
        if opt_soc["type"] == "cap":
            # On accumule ce que la société a réellement payé (plafond ou moins)
            # Pour les mois antérieurs on utilise le cap plein (pas de recalcul imbriqué)
            consumed[emp_name] += min(balance, opt_soc["val"])
        # Si type "rate" → le cap ne s'applique pas annuellement, pas d'accumulation

    return dict(consumed)


def process_data(lines, employees, params,
                 ytd_optique_consumed: dict = None,
                 ytd_emp_total: dict = None) -> list:
    """
    Retourne une liste de dicts enrichis.

    ytd_optique_consumed : {emp_name: fcfa_soc_payé_en_optique_mois_précédents}
      → applique le plafond optique annuel par employé de façon cumulée
    ytd_emp_total : {emp_name: fcfa_part_emp_ytd_précédents} (réservé future use)
    """
    rates  = params["rates"]
    prests = params["prestataires"]
    rows   = []

    for ln in lines:
        emp_id   = ln["x_studio_employee_inv"][0]
        emp_name = ln["x_studio_employee_inv"][1]
        emp_found = emp_id in employees
        emp_info  = employees.get(emp_id, {"name": emp_name, "dept": "", "client": "DI-Africa", "active": True})
        client    = emp_info["client"]

        # Alerte statut employé
        if not emp_found:
            emp_warning = f"Employé inconnu dans Odoo : {emp_name} (ID {emp_id}) — facture {ln['move_id'][1] if ln['move_id'] else '?'}"
        elif emp_info.get("active", True) is False:
            emp_warning = f"Employé archivé : {emp_name} — vérifier dette / retenue encore active"
        else:
            emp_warning = None

        partner_name = ln["partner_id"][1] if ln["partner_id"] else ""
        product_name = ln["product_id"][1] if ln["product_id"] else ""
        balance      = abs(float(ln["balance"] or 0))
        inv_ref      = ln["move_id"][1] if ln["move_id"] else ""

        # Type prestataire
        prest_type = prests.get(partner_name.upper(), None)
        if prest_type is None:
            prest_type = "Optique" if "optique" in product_name.lower() else "Pharmacie"

        # ── Classification article ────────────────────────────────────────
        article_type = _classify_article(product_name)

        # ── Calcul parts selon article_type ──────────────────────────────
        _RATE_100_EMP = {"consult_soc": 0.0, "consult_emp": 1.0,
                         "pharma_soc":  0.0, "pharma_emp":  1.0,
                         "optique_soc": None, "optique_emp": None,
                         "cc": "", "modele": "open bar", "plafond": None, "plafond_emp": None}
        _RATE_100_SOC = {"consult_soc": 1.0, "consult_emp": 0.0,
                         "pharma_soc":  1.0, "pharma_emp":  0.0,
                         "optique_soc": None, "optique_emp": None,
                         "cc": "", "modele": "open bar", "plafond": None, "plafond_emp": None}

        warning = None
        if article_type in ("accident", "insurance", "acpe_vm"):
            # 100 % employeur — pas de recherche dans les tables
            part_soc = balance
            part_emp = 0.0
            rate     = _RATE_100_SOC   # pour les champs cc/modele
        elif article_type == "ayants_droits_na":
            # 100 % employé (non autorisé)
            part_soc = 0.0
            part_emp = balance
            rate     = _RATE_100_EMP
        else:
            # "salarie" → table Taux Clients
            # "ayants_droits" → table Taux Ayants Droits (fallback → Taux Clients)
            rates_ad = params.get("rates_ad", {})
            if article_type == "ayants_droits" and rates_ad:
                rate = _find_rate(client, rates_ad)
                if rate is None:
                    # Pas de ligne AD pour ce client → fallback table principale avec warning
                    rate = _find_rate(client, rates)
                    if rate is not None:
                        warning = (f"Client '{client}' absent de 'Taux Ayants Droits' "
                                   f"→ taux Salarié appliqué par défaut")
            else:
                rate = _find_rate(client, rates)

            if rate is None:
                warning = f"Client '{client}' introuvable dans Paramètres → taux 50/50 appliqué"
                rate = {"consult_soc": 0.5, "consult_emp": 0.5,
                        "pharma_soc": 0.5,  "pharma_emp": 0.5,
                        "optique_soc": None, "optique_emp": None,
                        "cc": "", "modele": "open bar", "plafond": None, "plafond_emp": None}

            # Calcul selon type prestataire
            if prest_type == "Optique":
                opt_soc = rate["optique_soc"]
                opt_emp = rate["optique_emp"]
                if opt_soc is None:
                    part_soc = balance * 0.5
                    part_emp = balance * 0.5
                elif opt_soc.get("type") == "cap":
                    annual_cap = opt_soc["val"]
                    already    = (ytd_optique_consumed or {}).get(emp_name, 0.0)
                    remaining  = max(0.0, annual_cap - already)
                    part_soc   = min(balance, remaining)
                    part_emp   = max(0.0, balance - remaining)
                else:
                    soc_rate = opt_soc["val"]
                    emp_rate = opt_emp["val"] if opt_emp else (1.0 - soc_rate)
                    part_soc = balance * soc_rate
                    part_emp = balance * emp_rate
            elif prest_type == "Consultation":
                part_soc = balance * rate["consult_soc"]
                part_emp = balance * rate["consult_emp"]
            else:  # Pharmacie
                part_soc = balance * rate["pharma_soc"]
                part_emp = balance * rate["pharma_emp"]

        rows.append({
            "employee_name":  emp_name,
            "client":         client or "DI-Africa",
            "cc":             rate.get("cc", ""),
            "dept":           emp_info["dept"],
            "prestataire":    partner_name,
            "prest_type":     prest_type,
            "product":        product_name,
            "article_type":   article_type,   # "salarie"|"ayants_droits"|"ayants_droits_na"|"accident"|"insurance"|"acpe_vm"
            "date":           ln["date"],
            "invoice_ref":    inv_ref,
            "montant_total":  balance,
            "part_soc":       round(part_soc, 0),
            "part_emp":       round(part_emp, 0),
            "warning":          warning,
            "employee_warning": emp_warning,
            "modele":           rate.get("modele", "open bar"),
            "plafond_client": rate.get("plafond"),
            "plafond_emp":    rate.get("plafond_emp"),   # cap individuel annuel FCFA
        })

    return rows


# ── GÉNÉRATION EXCEL GLOBAL ───────────────────────────────────────────────
def _style_header(ws, row_num, cols, bg="1F4E79", fg="FFFFFF"):
    fill   = PatternFill("solid", fgColor=bg)
    font   = Font(bold=True, color=fg, name="Calibri", size=10)
    border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )
    for col in range(1, cols + 1):
        c = ws.cell(row=row_num, column=col)
        c.fill   = fill
        c.font   = font
        c.border = border
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def _style_data(ws, row_num, cols, alt=False):
    bg     = "F2F2F2" if alt else "FFFFFF"
    fill   = PatternFill("solid", fgColor=bg)
    border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )
    font = Font(name="Calibri", size=10)
    for col in range(1, cols + 1):
        c = ws.cell(row=row_num, column=col)
        c.fill   = fill
        c.border = border
        c.font   = font


def build_global_excel(rows: list, period_label: str,
                        rows_ytd: list = None, year: int = None) -> bytes:
    wb = openpyxl.Workbook()

    # ── Feuille Export (données brutes) ──────────────────────────────────
    ws = wb.active
    ws.title = "Export"
    headers = [
        "Date", "N° Facture", "Prestataire", "Type",
        "Employé(e)", "Client", "Produit",
        "Montant Total", "Part Société", "Part Employé(e)",
    ]
    ws.append(headers)
    _style_header(ws, 1, len(headers))

    for i, r in enumerate(rows):
        ws.append([
            r["date"], r["invoice_ref"], r["prestataire"], r["prest_type"],
            r["employee_name"], r["client"], r["product"],
            r["montant_total"], r["part_soc"], r["part_emp"],
        ])
        _style_data(ws, i + 2, len(headers), alt=(i % 2 == 1))
        # Format monétaire
        for col in [8, 9, 10]:
            ws.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'

    # Largeurs colonnes
    for col, w in zip(range(1, len(headers) + 1),
                      [12, 28, 32, 14, 32, 20, 22, 16, 16, 16]):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.freeze_panes = "A2"

    # ── Feuille TCD Clients ───────────────────────────────────────────────
    ws2 = wb.create_sheet("TCD Clients")
    hdrs = ["Client", "CC", "Consultation (Total)", "Pharmacie (Total)",
            "Optique (Total)", "Total", "Part Société", "Part Employé(e)"]
    ws2.append(hdrs)
    _style_header(ws2, 1, len(hdrs))

    from collections import defaultdict
    by_client = defaultdict(lambda: {"cc": "", "cons": 0, "phar": 0, "opti": 0,
                                      "soc": 0, "emp": 0})
    for r in rows:
        c = by_client[r["client"]]
        c["cc"] = r["cc"]
        if r["prest_type"] == "Consultation":
            c["cons"] += r["montant_total"]
        elif r["prest_type"] == "Optique":
            c["opti"] += r["montant_total"]
        else:
            c["phar"] += r["montant_total"]
        c["soc"] += r["part_soc"]
        c["emp"] += r["part_emp"]

    for i, (client, v) in enumerate(sorted(by_client.items())):
        total = v["cons"] + v["phar"] + v["opti"]
        ws2.append([client, v["cc"], v["cons"], v["phar"], v["opti"],
                    total, v["soc"], v["emp"]])
        _style_data(ws2, i + 2, len(hdrs), alt=(i % 2 == 1))
        for col in range(3, 9):
            ws2.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'

    for col, w in zip(range(1, len(hdrs) + 1),
                      [28, 12, 22, 18, 16, 16, 16, 16]):
        ws2.column_dimensions[get_column_letter(col)].width = w
    ws2.freeze_panes = "A2"

    # ── Feuille TCD Employés ──────────────────────────────────────────────
    ws3 = wb.create_sheet("TCD Employés")
    hdrs3 = ["Employé(e)", "Client", "Consultation", "Pharmacie",
             "Optique", "Total", "Part Société", "Part Employé(e)"]
    ws3.append(hdrs3)
    _style_header(ws3, 1, len(hdrs3))

    by_emp = defaultdict(lambda: {"client": "", "cons": 0, "phar": 0,
                                   "opti": 0, "soc": 0, "emp": 0})
    for r in rows:
        e = by_emp[r["employee_name"]]
        e["client"] = r["client"]
        if r["prest_type"] == "Consultation":
            e["cons"] += r["montant_total"]
        elif r["prest_type"] == "Optique":
            e["opti"] += r["montant_total"]
        else:
            e["phar"] += r["montant_total"]
        e["soc"] += r["part_soc"]
        e["emp"] += r["part_emp"]

    for i, (emp, v) in enumerate(sorted(by_emp.items())):
        total = v["cons"] + v["phar"] + v["opti"]
        ws3.append([emp, v["client"], v["cons"], v["phar"], v["opti"],
                    total, v["soc"], v["emp"]])
        _style_data(ws3, i + 2, len(hdrs3), alt=(i % 2 == 1))
        for col in range(3, 9):
            ws3.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'

    for col, w in zip(range(1, len(hdrs3) + 1),
                      [34, 28, 18, 16, 14, 16, 16, 16]):
        ws3.column_dimensions[get_column_letter(col)].width = w
    ws3.freeze_panes = "A2"

    # ── Feuille YTD Clients (si données disponibles) ──────────────────────
    if rows_ytd:
        ws4 = wb.create_sheet("YTD Clients")
        yr_label = str(year) if year else ""
        hdrs4 = ["Client", "CC", "Modèle", "Plafond Annuel",
                 "Consul. YTD", "Pharma. YTD", "Optique YTD",
                 "Total YTD", "Part Soc. YTD", "Part Emp. YTD",
                 "% Plafond consommé"]
        ws4.append(hdrs4)
        _style_header(ws4, 1, len(hdrs4))

        from collections import defaultdict as _dd
        ytd_by_client = _dd(lambda: {"cc": "", "modele": "", "plafond": None,
                                      "cons": 0, "phar": 0, "opti": 0,
                                      "soc": 0, "emp": 0})
        for r in rows_ytd:
            c = ytd_by_client[r["client"]]
            c["cc"]     = r.get("cc", "")
            c["modele"] = r.get("modele", "open bar")
            c["plafond"]= r.get("plafond_client")
            if r["prest_type"] == "Consultation":
                c["cons"] += r["montant_total"]
            elif r["prest_type"] == "Optique":
                c["opti"] += r["montant_total"]
            else:
                c["phar"] += r["montant_total"]
            c["soc"] += r["part_soc"]
            c["emp"] += r["part_emp"]

        for i, (client, v) in enumerate(sorted(ytd_by_client.items())):
            total = v["cons"] + v["phar"] + v["opti"]
            plafond = v["plafond"]
            pct_plafond = (total / plafond * 100) if plafond else None
            ws4.append([
                client, v["cc"],
                "Provision" if v["modele"] == "provision" else "Open bar",
                plafond if plafond else "—",
                v["cons"], v["phar"], v["opti"],
                total, v["soc"], v["emp"],
                f"{pct_plafond:.1f} %" if pct_plafond is not None else "—",
            ])
            row_n = i + 2
            _style_data(ws4, row_n, len(hdrs4), alt=(i % 2 == 1))
            for col in [4, 5, 6, 7, 8, 9, 10]:
                cell = ws4.cell(row=row_n, column=col)
                if isinstance(cell.value, (int, float)):
                    cell.number_format = '#,##0" FCFA"'
            # Colorier en orange si dépassement plafond
            if pct_plafond is not None and pct_plafond >= 100:
                for col in range(1, len(hdrs4) + 1):
                    ws4.cell(row=row_n, column=col).fill = PatternFill("solid", fgColor="FFE0CC")
            elif pct_plafond is not None and pct_plafond >= 80:
                for col in range(1, len(hdrs4) + 1):
                    ws4.cell(row=row_n, column=col).fill = PatternFill("solid", fgColor="FFF2CC")

        for col, w in zip(range(1, len(hdrs4) + 1),
                          [28, 10, 12, 18, 18, 16, 14, 16, 16, 16, 18]):
            ws4.column_dimensions[get_column_letter(col)].width = w
        ws4.freeze_panes = "A2"

        # ── Feuille YTD Employés ───────────────────────────────────────────
        ws5 = wb.create_sheet("YTD Employés")
        hdrs5 = ["Employé(e)", "Client", "Consul. YTD", "Pharma. YTD",
                 "Optique YTD", "Total YTD", "Part Soc. YTD", "Part Emp. YTD"]
        ws5.append(hdrs5)
        _style_header(ws5, 1, len(hdrs5))

        ytd_by_emp = _dd(lambda: {"client": "", "cons": 0, "phar": 0,
                                   "opti": 0, "soc": 0, "emp": 0})
        for r in rows_ytd:
            e = ytd_by_emp[r["employee_name"]]
            e["client"] = r["client"]
            if r["prest_type"] == "Consultation":
                e["cons"] += r["montant_total"]
            elif r["prest_type"] == "Optique":
                e["opti"] += r["montant_total"]
            else:
                e["phar"] += r["montant_total"]
            e["soc"] += r["part_soc"]
            e["emp"] += r["part_emp"]

        for i, (emp, v) in enumerate(sorted(ytd_by_emp.items())):
            total = v["cons"] + v["phar"] + v["opti"]
            ws5.append([emp, v["client"], v["cons"], v["phar"], v["opti"],
                        total, v["soc"], v["emp"]])
            _style_data(ws5, i + 2, len(hdrs5), alt=(i % 2 == 1))
            for col in range(3, 9):
                ws5.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'

        for col, w in zip(range(1, len(hdrs5) + 1),
                          [34, 28, 18, 16, 14, 16, 16, 16]):
            ws5.column_dimensions[get_column_letter(col)].width = w
        ws5.freeze_panes = "A2"

        # ── Feuille Top Consommateurs ──────────────────────────────────────
        ws6 = wb.create_sheet("Top Consommateurs")
        yr_label6 = str(year) if year else ""

        # Titre
        ws6.merge_cells("A1:J1")
        ws6["A1"] = f"Top Consommateurs Santé — {yr_label6} (Jan → {period_label})"
        ws6["A1"].font      = Font(bold=True, size=13, color="FFFFFF", name="Calibri")
        ws6["A1"].fill      = PatternFill("solid", fgColor="1F4E79")
        ws6["A1"].alignment = Alignment(horizontal="center")
        ws6.row_dimensions[1].height = 24

        hdrs6 = ["#", "Employé(e)", "Client", "Statut",
                 "Total YTD", "Part Soc. YTD", "Part Emp. YTD",
                 "Consul. YTD", "Pharma. YTD", "Optique YTD"]
        ws6.append(hdrs6)
        _style_header(ws6, 2, len(hdrs6))

        # Réutiliser ytd_by_emp (déjà construit pour ws5)
        # + enrichir avec statut actif/archivé si disponible depuis rows_ytd
        emp_warnings_map = {}
        for r in rows_ytd:
            w = r.get("employee_warning")
            if w:
                emp_warnings_map[r["employee_name"]] = w

        ranked = sorted(ytd_by_emp.items(),
                        key=lambda x: x[1]["cons"] + x[1]["phar"] + x[1]["opti"],
                        reverse=True)

        # Paliers couleur (fonds) : top 10% rouge, top 30% orange, reste normal
        n_emp = len(ranked)
        fill_red    = PatternFill("solid", fgColor="FFDAD9")
        fill_orange = PatternFill("solid", fgColor="FFF2CC")
        fill_arch   = PatternFill("solid", fgColor="E2EFDA")  # vert pâle = archivé/inconnu

        for i, (emp, v) in enumerate(ranked):
            total = v["cons"] + v["phar"] + v["opti"]
            rank  = i + 1

            # Statut employé
            emp_w = emp_warnings_map.get(emp, "")
            if "archivé" in emp_w:
                statut = "⚠️ Archivé"
            elif "inconnu" in emp_w:
                statut = "❓ Inconnu"
            else:
                statut = "✅ Actif"

            ws6.append([rank, emp, v["client"], statut,
                        total, v["soc"], v["emp"],
                        v["cons"], v["phar"], v["opti"]])

            data_row = i + 3
            # Couleur de fond selon rang / statut
            if emp_w:
                row_fill = fill_arch
            elif n_emp >= 10 and rank <= max(1, n_emp // 10):
                row_fill = fill_red
            elif rank <= max(3, n_emp // 3):
                row_fill = fill_orange
            else:
                row_fill = PatternFill("solid", fgColor="F9F9F9") if i % 2 == 1 else None

            for col_idx in range(1, len(hdrs6) + 1):
                cell = ws6.cell(row=data_row, column=col_idx)
                if row_fill:
                    cell.fill = row_fill
                cell.font      = Font(name="Calibri", size=10)
                cell.alignment = Alignment(horizontal="left")

            for col_idx in range(5, 11):  # colonnes montants
                ws6.cell(row=data_row, column=col_idx).number_format = '#,##0" FCFA"'
            ws6.cell(row=data_row, column=1).alignment = Alignment(horizontal="center")

        # Ligne totaux
        last_row6 = 3 + len(ranked)
        grand_total = sum(v["cons"] + v["phar"] + v["opti"] for _, v in ytd_by_emp.items())
        grand_soc   = sum(v["soc"] for _, v in ytd_by_emp.items())
        grand_emp   = sum(v["emp"] for _, v in ytd_by_emp.items())
        grand_cons  = sum(v["cons"] for _, v in ytd_by_emp.items())
        grand_phar  = sum(v["phar"] for _, v in ytd_by_emp.items())
        grand_opti  = sum(v["opti"] for _, v in ytd_by_emp.items())
        ws6.append(["", f"TOTAL ({len(ranked)} employés)", "", "",
                    grand_total, grand_soc, grand_emp,
                    grand_cons, grand_phar, grand_opti])
        _style_header(ws6, last_row6, len(hdrs6), bg="2E75B6")
        for col_idx in range(5, 11):
            ws6.cell(row=last_row6, column=col_idx).number_format = '#,##0" FCFA"'

        for col, w in zip(range(1, len(hdrs6) + 1),
                          [5, 34, 22, 12, 18, 16, 16, 16, 14, 14]):
            ws6.column_dimensions[get_column_letter(col)].width = w
        ws6.freeze_panes = "B3"

    # ── Titre et métadonnées ──────────────────────────────────────────────
    for sheet in wb.sheetnames:
        wb[sheet].sheet_properties.tabColor = "1F4E79"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── GÉNÉRATION EXCEL INDIVIDUEL ───────────────────────────────────────────
def build_individual_excel(emp_name: str, emp_rows: list, period_label: str) -> bytes:
    wb  = openpyxl.Workbook()
    ws  = wb.active
    ws.title = emp_name[:31]  # max 31 chars pour onglet

    # En-tête
    ws.merge_cells("A1:J1")
    ws["A1"] = f"Relevé de consommation santé — {emp_name} — {period_label}"
    ws["A1"].font      = Font(bold=True, size=13, color="1F4E79", name="Calibri")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[1].height = 24

    headers = ["Date", "N° Facture", "Prestataire", "Type",
               "Produit", "Montant Total", "Part Société", "Part Employé(e)"]
    ws.append(headers)
    _style_header(ws, 2, len(headers))

    total_soc = total_emp = 0
    for i, r in enumerate(emp_rows):
        ws.append([
            r["date"], r["invoice_ref"], r["prestataire"], r["prest_type"],
            r["product"], r["montant_total"], r["part_soc"], r["part_emp"],
        ])
        _style_data(ws, i + 3, len(headers), alt=(i % 2 == 1))
        for col in [6, 7, 8]:
            ws.cell(row=i + 3, column=col).number_format = '#,##0" FCFA"'
        total_soc += r["part_soc"]
        total_emp += r["part_emp"]

    # Ligne totaux
    last_data = 3 + len(emp_rows)
    ws.append(["", "", "", "", "TOTAL",
               total_soc + total_emp, total_soc, total_emp])
    _style_header(ws, last_data, len(headers), bg="2E75B6")
    for col in [6, 7, 8]:
        ws.cell(row=last_data, column=col).number_format = '#,##0" FCFA"'

    for col, w in zip(range(1, len(headers) + 1),
                      [12, 28, 32, 14, 22, 18, 16, 16]):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.freeze_panes = "A3"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── GÉNÉRATION RAPPORT CLIENT (REFACTURATION) ─────────────────────────────
def _taux_label(row: dict) -> str:
    """Affiche le taux/plafond en texte lisible pour le rapport client."""
    total = row["montant_total"]
    if total <= 0:
        return "—"
    if row["prest_type"] == "Optique":
        if row["part_emp"] == 0:
            return "100 % soc"
        elif row["part_soc"] >= total:
            return "100 % soc"
        else:
            pct = round(row["part_soc"] / total * 100)
            return f"Plafond / {pct} % soc" if pct < 100 else "100 % soc"
    pct = round(row["part_soc"] / total * 100)
    return f"{pct} %"


def build_client_excel(client_name: str, client_rows: list, period_label: str) -> bytes:
    """
    Rapport de consommation par client — ce que DI-Africa refacture au client.
    Format identique aux anciens exports GAS :
      Date | N° Facture | Employé | Prestataire | Type | Montant Total
      | Taux/Plafond | Part Employé(e) | Part Employeur (Société)
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = client_name[:31]

    # ── En-tête titre ──────────────────────────────────────────────────────
    NCOLS = 9
    ws.merge_cells(f"A1:{get_column_letter(NCOLS)}1")
    ws["A1"] = f"Rapport de consommation — {client_name} — {period_label}"
    ws["A1"].font      = Font(bold=True, size=13, color="FFFFFF", name="Calibri")
    ws["A1"].fill      = PatternFill("solid", fgColor="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells(f"A2:{get_column_letter(NCOLS)}2")
    ws["A2"] = "DI-Africa (Congo) SA · Direction Médicale · sante@di-africa.com"
    ws["A2"].font      = Font(italic=True, size=9, color="595959", name="Calibri")
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 16

    # ── En-têtes colonnes ──────────────────────────────────────────────────
    headers = [
        "Date", "N° Facture", "Employé(e)", "Prestataire",
        "Type", "Montant Total", "Taux / Plafond",
        "Part Employé(e)", "Part Employeur",
    ]
    ws.append(headers)
    _style_header(ws, 3, NCOLS)

    # ── Données ────────────────────────────────────────────────────────────
    total_montant = total_emp = total_soc = 0
    for i, r in enumerate(client_rows):
        ws.append([
            r["date"],
            r["invoice_ref"],
            r["employee_name"],
            r["prestataire"],
            r["prest_type"],
            r["montant_total"],
            _taux_label(r),
            r["part_emp"],
            r["part_soc"],
        ])
        row_num = i + 4
        _style_data(ws, row_num, NCOLS, alt=(i % 2 == 1))
        for col in [6, 8, 9]:
            ws.cell(row=row_num, column=col).number_format = '#,##0" FCFA"'
        total_montant += r["montant_total"]
        total_emp     += r["part_emp"]
        total_soc     += r["part_soc"]

    # ── Ligne TOTAL ────────────────────────────────────────────────────────
    last = 4 + len(client_rows)
    ws.append(["", "", "", "", "TOTAL",
               total_montant, "", total_emp, total_soc])
    _style_header(ws, last, NCOLS, bg="2E75B6")
    for col in [6, 8, 9]:
        ws.cell(row=last, column=col).number_format = '#,##0" FCFA"'

    # ── Largeurs ───────────────────────────────────────────────────────────
    for col, w in zip(range(1, NCOLS + 1),
                      [12, 26, 34, 32, 14, 18, 16, 18, 18]):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.freeze_panes = "A4"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_clients_recap_excel(rows: list, period_label: str) -> bytes:
    """
    Récapitulatif multi-onglets : un onglet par client + un onglet synthèse.
    Utile pour le reporting global de refacturation.
    """
    wb = openpyxl.Workbook()

    # ── Onglet Synthèse ────────────────────────────────────────────────────
    ws_synth = wb.active
    ws_synth.title = "Synthèse"
    hdrs = ["Client", "CC", "Nb employés", "Consultation",
            "Pharmacie", "Optique", "Total", "Part Employé(e)", "Part Employeur"]
    ws_synth.append(hdrs)
    _style_header(ws_synth, 1, len(hdrs))

    by_client = defaultdict(lambda: {
        "cc": "", "emps": set(),
        "cons": 0, "phar": 0, "opti": 0, "emp": 0, "soc": 0
    })
    for r in rows:
        c = by_client[r["client"]]
        c["cc"] = r["cc"]
        c["emps"].add(r["employee_name"])
        if r["prest_type"] == "Consultation":
            c["cons"] += r["montant_total"]
        elif r["prest_type"] == "Optique":
            c["opti"] += r["montant_total"]
        else:
            c["phar"] += r["montant_total"]
        c["emp"] += r["part_emp"]
        c["soc"] += r["part_soc"]

    for i, (client, v) in enumerate(sorted(by_client.items())):
        total = v["cons"] + v["phar"] + v["opti"]
        ws_synth.append([client, v["cc"], len(v["emps"]),
                         v["cons"], v["phar"], v["opti"],
                         total, v["emp"], v["soc"]])
        _style_data(ws_synth, i + 2, len(hdrs), alt=(i % 2 == 1))
        for col in [4, 5, 6, 7, 8, 9]:
            ws_synth.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'

    for col, w in zip(range(1, len(hdrs) + 1),
                      [32, 10, 14, 18, 16, 14, 16, 18, 18]):
        ws_synth.column_dimensions[get_column_letter(col)].width = w
    ws_synth.freeze_panes = "A2"

    # ── Un onglet par client ───────────────────────────────────────────────
    by_client_rows = defaultdict(list)
    for r in rows:
        by_client_rows[r["client"]].append(r)

    for client, c_rows in sorted(by_client_rows.items()):
        ws = wb.create_sheet(title=client[:31])
        headers = ["Date", "N° Facture", "Employé(e)", "Prestataire",
                   "Type", "Montant Total", "Taux / Plafond",
                   "Part Employé(e)", "Part Employeur"]
        ws.append(headers)
        _style_header(ws, 1, len(headers))
        tot_m = tot_e = tot_s = 0
        for i, r in enumerate(c_rows):
            ws.append([r["date"], r["invoice_ref"], r["employee_name"],
                       r["prestataire"], r["prest_type"], r["montant_total"],
                       _taux_label(r), r["part_emp"], r["part_soc"]])
            _style_data(ws, i + 2, len(headers), alt=(i % 2 == 1))
            for col in [6, 8, 9]:
                ws.cell(row=i + 2, column=col).number_format = '#,##0" FCFA"'
            tot_m += r["montant_total"]
            tot_e += r["part_emp"]
            tot_s += r["part_soc"]
        last = 2 + len(c_rows)
        ws.append(["", "", "", "", "TOTAL", tot_m, "", tot_e, tot_s])
        _style_header(ws, last, len(headers), bg="2E75B6")
        for col in [6, 8, 9]:
            ws.cell(row=last, column=col).number_format = '#,##0" FCFA"'
        for col, w in zip(range(1, len(headers) + 1),
                          [12, 26, 34, 32, 14, 18, 16, 18, 18]):
            ws.column_dimensions[get_column_letter(col)].width = w
        ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── RAPPORT PROVISION SANTÉ ──────────────────────────────────────────────
def build_provision_excel(rows_ytd: list, params: dict,
                           period_label: str, year: int) -> bytes:
    """
    Rapport de suivi des provisions santé par client :
    - Plafond annuel configuré
    - Consommé YTD (part société)
    - Restant disponible
    - Alertes dépassement / approche plafond
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Suivi Provision"

    # ── Titre ──────────────────────────────────────────────────────────────
    NCOLS = 10
    ws.merge_cells(f"A1:{get_column_letter(NCOLS)}1")
    ws["A1"] = f"Suivi Provision Santé — {year} — Cumulé au {period_label}"
    ws["A1"].font      = Font(bold=True, size=13, color="FFFFFF", name="Calibri")
    ws["A1"].fill      = PatternFill("solid", fgColor="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells(f"A2:{get_column_letter(NCOLS)}2")
    ws["A2"] = "DI-Africa (Congo) SA · Direction Médicale · sante@di-africa.com"
    ws["A2"].font      = Font(italic=True, size=9, color="595959", name="Calibri")
    ws["A2"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[2].height = 16

    # ── En-têtes ───────────────────────────────────────────────────────────
    headers = [
        "Client", "CC", "Modèle",
        "Plafond Annuel", "Conso. Totale YTD", "Part Soc. YTD", "Part Emp. YTD",
        "Restant Disponible", "% Consommé", "Statut",
    ]
    ws.append(headers)
    _style_header(ws, 3, NCOLS)

    # ── Agrégation YTD par client ──────────────────────────────────────────
    by_client = defaultdict(lambda: {"cc": "", "modele": "open bar", "plafond": None,
                                      "total": 0, "soc": 0, "emp": 0})
    rates = params.get("rates", {})
    for r in rows_ytd:
        c = by_client[r["client"]]
        c["total"] += r["montant_total"]
        c["soc"]   += r["part_soc"]
        c["emp"]   += r["part_emp"]
        # Récupérer modèle et plafond depuis params
        rate_info = rates.get(r["client"])
        if rate_info is None:
            for k in rates:
                if k.lower() == r["client"].lower():
                    rate_info = rates[k]
                    break
        if rate_info:
            c["cc"]      = rate_info.get("cc", "")
            c["modele"]  = rate_info.get("modele", "open bar")
            c["plafond"] = rate_info.get("plafond")

    # ── Lignes ─────────────────────────────────────────────────────────────
    ORANGE = "FFE0CC"  # dépassement
    YELLOW = "FFF2CC"  # approche (≥80%)
    GREEN  = "E2EFDA"  # ok

    for i, (client, v) in enumerate(sorted(by_client.items())):
        plafond = v["plafond"]
        modele  = v["modele"]

        if plafond and modele == "provision":
            restant    = max(0, plafond - v["soc"])
            pct        = v["soc"] / plafond * 100
            pct_txt    = f"{pct:.1f} %"
            restant_txt = restant
            if pct >= 100:
                statut = "🔴 Dépassé"
                bg = ORANGE
            elif pct >= 80:
                statut = "🟡 Attention"
                bg = YELLOW
            else:
                statut = "🟢 OK"
                bg = GREEN
        else:
            restant_txt = "—"
            pct_txt     = "—"
            statut      = "Open bar"
            bg          = "FFFFFF"

        ws.append([
            client,
            v["cc"],
            "Provision" if modele == "provision" else "Open bar",
            plafond if plafond else "—",
            v["total"],
            v["soc"],
            v["emp"],
            restant_txt,
            pct_txt,
            statut,
        ])
        row_n = i + 4
        _style_data(ws, row_n, NCOLS)
        if bg != "FFFFFF":
            for col in range(1, NCOLS + 1):
                ws.cell(row=row_n, column=col).fill = PatternFill("solid", fgColor=bg)
        for col in [4, 5, 6, 7]:
            cell = ws.cell(row=row_n, column=col)
            if isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0" FCFA"'
        if isinstance(restant_txt, (int, float)):
            ws.cell(row=row_n, column=8).number_format = '#,##0" FCFA"'

    # Ligne totaux
    last = 4 + len(by_client)
    total_conso  = sum(v["total"] for v in by_client.values())
    total_soc    = sum(v["soc"]   for v in by_client.values())
    total_emp    = sum(v["emp"]   for v in by_client.values())
    total_plafond = sum(v["plafond"] for v in by_client.values()
                        if v["plafond"] and v["modele"] == "provision") or None
    ws.append(["", "", "TOTAL",
                total_plafond if total_plafond else "—",
                total_conso, total_soc, total_emp,
                "—", "—", ""])
    _style_header(ws, last, NCOLS, bg="2E75B6")
    for col in [5, 6, 7]:
        ws.cell(row=last, column=col).number_format = '#,##0" FCFA"'
    if isinstance(total_plafond, (int, float)):
        ws.cell(row=last, column=4).number_format = '#,##0" FCFA"'

    # ── Largeurs ───────────────────────────────────────────────────────────
    for col, w in zip(range(1, NCOLS + 1),
                      [30, 10, 12, 20, 20, 18, 18, 20, 14, 14]):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.freeze_panes = "A4"

    # ── Légende ────────────────────────────────────────────────────────────
    ws_leg = wb.create_sheet("Légende")
    ws_leg["A1"] = "Légende — Statuts Provision Santé"
    ws_leg["A1"].font = Font(bold=True, size=11, name="Calibri")
    legends = [
        ("🔴 Dépassé",   "Consommation ≥ 100% du plafond",  ORANGE),
        ("🟡 Attention", "Consommation entre 80% et 99% du plafond", YELLOW),
        ("🟢 OK",        "Consommation < 80% du plafond",    GREEN),
        ("Open bar",    "Pas de plafond configuré — refacturation intégrale", "FFFFFF"),
    ]
    for j, (statut, desc, color) in enumerate(legends, start=2):
        ws_leg[f"A{j}"] = statut
        ws_leg[f"B{j}"] = desc
        ws_leg[f"A{j}"].fill = PatternFill("solid", fgColor=color)
        ws_leg[f"B{j}"].fill = PatternFill("solid", fgColor=color)
    ws_leg.column_dimensions["A"].width = 18
    ws_leg.column_dimensions["B"].width = 55

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── GÉNÉRATION RELEVÉ AUTOMATIQUE PAR EMPLOYÉ (DOCX) ─────────────────────
def build_releve_employe_docx(emp_name: str, client: str,
                               emp_rows: list, period_label: str) -> bytes:
    """
    Relevé de consommation individuel généré automatiquement (sans template).
    Format identique aux anciens exports : en-tête DI-Africa, tableau,
    colonne Taux/Plafond, ligne TOTAL.
    """
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Marges ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    def _para(text="", bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_cell_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{side}")
                    border.set(qn("w:val"),   "single")
                    border.set(qn("w:sz"),    "4")
                    border.set(qn("w:space"), "0")
                    border.set(qn("w:color"), "BFBFBF")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

    # ── En-tête société ──────────────────────────────────────────────────────
    _para("DI Africa (Congo) SA", bold=True, size=13,
          color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    _para("Pointe-Noire, Congo  ·  sante@di-africa.com",
          size=9, color=(89, 89, 89), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _para("RELEVÉ DE CONSOMMATION SANTÉ", bold=True, size=14,
          color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(f"{emp_name}  ·  {client}  ·  {period_label}",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ── Tableau ──────────────────────────────────────────────────────────────
    col_headers = ["Date", "N° Facture", "Prestataire", "Type",
                   "Montant Total", "Taux / Plafond", "Part Employé(e)", "Part Employeur"]
    col_widths  = [Cm(2.2), Cm(3.4), Cm(4.2), Cm(2.8),
                   Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8)]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = None  # pas de style nommé pour éviter les conflits

    # Largeurs colonnes
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Ligne en-tête
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(hdr_cells[i], "1F4E79")

    # Lignes données
    total_montant = total_emp = total_soc = 0
    for idx, r in enumerate(emp_rows):
        row_cells = table.add_row().cells
        bg = "F2F2F2" if idx % 2 == 1 else "FFFFFF"
        values = [
            str(r["date"]),
            r["invoice_ref"],
            r["prestataire"],
            r["prest_type"],
            f"{r['montant_total']:,.0f}",
            _taux_label(r),
            f"{r['part_emp']:,.0f}",
            f"{r['part_soc']:,.0f}",
        ]
        for i, v in enumerate(values):
            row_cells[i].text = v
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run(v)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            _set_cell_bg(row_cells[i], bg)
        total_montant += r["montant_total"]
        total_emp     += r["part_emp"]
        total_soc     += r["part_soc"]

    # Ligne TOTAL
    tot_cells = table.add_row().cells
    tot_vals  = ["", "", "", "TOTAL",
                 f"{total_montant:,.0f}", "",
                 f"{total_emp:,.0f}", f"{total_soc:,.0f}"]
    for i, v in enumerate(tot_vals):
        tot_cells[i].text = v
        p = tot_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0] if p.runs else p.add_run(v)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(tot_cells[i], "2E75B6")

    _set_cell_borders(table)

    # ── Pied de page ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    _para(
        "S.A au capital de 20 000 000 FCFA · Rue Germain Bicoumat, Pointe-Noire, CONGO\n"
        "N° RCCM CG-PNR-01-2021-B14-00003 · NIU M210000001927504 · 04 063 6397",
        size=7, color=(127, 127, 127), align=WD_ALIGN_PARAGRAPH.CENTER
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── GÉNÉRATION BON DE CONSOMMATION (DOCX) ─────────────────────────────────
def fill_bon_template(template_bytes: bytes, emp_name: str,
                      emp_rows: list, period_label: str,
                      dest_type: str) -> bytes:
    """
    Remplace <<NOM>> et <<TABLEAU_CONSOMMATIONS>> dans le template .docx.
    dest_type = 'Employeur' ou 'Employé'
    """
    doc = Document(io.BytesIO(template_bytes))

    # Remplacement <<NOM>>
    for para in doc.paragraphs:
        if "<<NOM>>" in para.text:
            for run in para.runs:
                run.text = run.text.replace("<<NOM>>", emp_name)
        if "<<PERIODE>>" in para.text:
            for run in para.runs:
                run.text = run.text.replace("<<PERIODE>>", period_label)

    # Remplacement <<TABLEAU_CONSOMMATIONS>>
    for para in doc.paragraphs:
        if "<<TABLEAU_CONSOMMATIONS>>" in para.text:
            # Vider le paragraphe placeholder
            para.clear()
            # Corriger les marges si elles sont stockées en float dans le template
            # (bug courant avec certains templates Word → ValueError in python-docx)
            # On patch directement les attributs XML pour convertir float → int
            import re as _re
            for section in doc.sections:
                pgMar = section._sectPr.find(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgMar"
                )
                if pgMar is not None:
                    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    for attr in ["left", "right", "top", "bottom", "header", "footer", "gutter"]:
                        key = f"{{{ns}}}{attr}"
                        val = pgMar.get(key)
                        if val and "." in val:
                            pgMar.set(key, str(int(float(val))))
            # Construire un tableau dans le document
            table = doc.add_table(rows=1, cols=5)
            # Ne pas forcer un style nommé — le template peut ne pas l'inclure
            # On appliquera des bordures manuellement si nécessaire
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(["Date", "Prestataire", "Montant", "Part Société", "Part Employé(e)"]):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].bold = True

            total_soc = total_emp = 0
            for r in emp_rows:
                row_cells = table.add_row().cells
                row_cells[0].text = str(r["date"])
                row_cells[1].text = r["prestataire"]
                row_cells[2].text = f"{r['montant_total']:,.0f} FCFA"
                row_cells[3].text = f"{r['part_soc']:,.0f} FCFA"
                row_cells[4].text = f"{r['part_emp']:,.0f} FCFA"
                total_soc += r["part_soc"]
                total_emp += r["part_emp"]

            # Ligne totaux
            tot_cells = table.add_row().cells
            tot_cells[0].text = "TOTAL"
            tot_cells[0].paragraphs[0].runs[0].bold = True
            tot_cells[2].text = f"{total_soc + total_emp:,.0f} FCFA"
            tot_cells[3].text = f"{total_soc:,.0f} FCFA"
            tot_cells[4].text = f"{total_emp:,.0f} FCFA"

            # Déplacer le tableau avant le paragraphe placeholder
            para._element.addprevious(table._element)
            break

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── INTERFACE STREAMLIT ───────────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="AutoSanté — DI-Africa Congo",
        page_icon="🏥",
        layout="wide",
    )

    st.title("🏥 AutoSanté — DI-Africa (Congo) SA")
    st.caption("Génération des rapports de consommation santé · Version Odoo API direct")

    # ── Connexion Odoo ────────────────────────────────────────────────────
    uid, models = odoo_connect()
    st.success(f"✅ Connecté à Odoo (`{ODOO_URL}`) · Société : DI-Africa (Congo) SA")

    st.divider()

    # ── Paramètres depuis Google Sheets (automatique) ─────────────────────
    col_gs1, col_gs2 = st.columns([3, 1])
    with col_gs1:
        st.info(
            "📊 Les taux clients et prestataires sont chargés automatiquement "
            "depuis le **Google Sheet Paramètres** partagé. "
            "Aucun fichier à uploader."
        )
    with col_gs2:
        sheet_url = (f"https://docs.google.com/spreadsheets/d/{PARAMS_SHEET_ID}/edit")
        st.markdown(f"[🔗 Ouvrir le Google Sheet]({sheet_url})", unsafe_allow_html=False)
        if st.button("🔄 Recharger les paramètres"):
            load_params_from_gsheet.clear()
            st.rerun()

    # ── Sélection période ─────────────────────────────────────────────────
    col3, col4 = st.columns(2)
    with col3:
        month = st.selectbox("Mois", range(1, 13),
                             format_func=lambda m: datetime(2000, m, 1).strftime("%B").capitalize(),
                             index=datetime.today().month - 2 if datetime.today().month > 1 else 0)
    with col4:
        year  = st.number_input("Année", min_value=2020,
                                max_value=2030, value=datetime.today().year)

    last_day   = calendar.monthrange(year, month)[1]
    date_from  = f"{year}-{month:02d}-01"
    date_to    = f"{year}-{month:02d}-{last_day:02d}"
    period_label = f"{datetime(year, month, 1).strftime('%B %Y').capitalize()}"
    st.caption(f"Période : **{date_from}** → **{date_to}**")

    st.divider()

    # ── Templates Bons de consommation (optionnel) ────────────────────────
    with st.expander("📄 Templates Bon de consommation (optionnel)"):
        tpl_emp    = st.file_uploader("Template Employé",   type=["docx"], key="tpl_emp")
        tpl_empeur = st.file_uploader("Template Employeur", type=["docx"], key="tpl_empeur")

    st.divider()

    # ── Option YTD ────────────────────────────────────────────────────────
    include_ytd = st.checkbox(
        "📅 Inclure les données cumulées YTD (Jan → mois sélectionné)",
        value=True,
        help="Ajoute les onglets 'YTD Clients' et 'YTD Employés' dans le rapport global "
             "et génère le rapport de suivi provision santé."
    )

    st.divider()

    # ── Bouton de génération ──────────────────────────────────────────────
    if st.button("🚀 Générer les rapports", type="primary", use_container_width=True):

        with st.spinner("Récupération des données Odoo…"):
            lines     = fetch_invoice_lines(uid, models, date_from, date_to)
            employees = fetch_employees(uid, models)

        st.info(f"📊 {len(lines)} lignes de facturation récupérées pour {period_label}")

        if not lines:
            st.warning("Aucune ligne trouvée pour cette période. Vérifiez les dates.")
            st.stop()

        with st.spinner("Chargement des paramètres…"):
            params = load_params_from_gsheet(PARAMS_SHEET_ID)
        st.success(f"✅ Paramètres chargés : "
                   f"{len(params['rates'])} clients · "
                   f"{len(params['prestataires'])} prestataires")

        # ── Données YTD + calcul optique cumulatif ───────────────────────────
        rows_ytd   = None
        ytd_label  = period_label
        ytd_optique_consumed = None

        if include_ytd:
            with st.spinner("Récupération des données YTD (Jan → mois sélectionné)…"):
                lines_ytd = fetch_invoice_lines_ytd(uid, models, year, month)

            if month > 1:
                # Séparer les mois précédents (Jan→M-1) pour le calcul optique cumulatif
                lines_prior = [l for l in lines_ytd if l["date"] < date_from]
                with st.spinner("Calcul cumul optique YTD…"):
                    ytd_optique_consumed = compute_ytd_optique_consumed(
                        lines_prior, employees, params
                    )
                ytd_label = f"Jan → {period_label}"
                st.info(f"📅 YTD : {len(lines_ytd)} lignes sur {ytd_label}")

        with st.spinner("Calcul des parts société / employé(e)…"):
            # Le calcul mensuel tient compte du cap optique déjà consommé YTD
            rows = process_data(lines, employees, params,
                                ytd_optique_consumed=ytd_optique_consumed)
            # Mémoriser pour Phase 2 (retenues)
            st.session_state["rows"]         = rows
            st.session_state["params"]       = params
            st.session_state["period_label"] = period_label
            st.session_state["year"]         = year
            st.session_state["month"]        = month

        if include_ytd:
            with st.spinner("Calcul YTD complet…"):
                if month > 1:
                    rows_ytd = process_data(lines_ytd, employees, params,
                                            ytd_optique_consumed=None)
                    # Note: pour le YTD global, on recalcule sans correction (agrégat)
                else:
                    rows_ytd = rows  # Janvier : YTD = mois

        # Avertissements client / taux
        warnings = [r for r in rows if r["warning"]]
        if warnings:
            with st.expander(f"⚠️ {len(warnings)} avertissement(s) de matching"):
                for w in set(r["warning"] for r in warnings):
                    st.warning(w)

        # ── Alertes Employés inconnus / archivés ──────────────────────────
        emp_warnings = [r for r in rows if r.get("employee_warning")]
        if emp_warnings:
            # Dédoublonner par message + regrouper les montants
            from collections import defaultdict as _dd
            grouped = _dd(float)
            for r in emp_warnings:
                grouped[r["employee_warning"]] += r["montant_total"]
            archived = {k: v for k, v in grouped.items() if k.startswith("Employé archivé")}
            unknown  = {k: v for k, v in grouped.items() if k.startswith("Employé inconnu")}
            label = []
            if unknown:  label.append(f"{len(unknown)} inconnu(s)")
            if archived: label.append(f"{len(archived)} archivé(s)")
            with st.expander(f"🔴 Employés à vérifier : {' · '.join(label)}", expanded=True):
                if unknown:
                    st.error("**Employés non trouvés dans Odoo** — factures potentiellement mal rattachées :")
                    for msg, total in sorted(unknown.items(), key=lambda x: -x[1]):
                        st.markdown(f"- {msg} · **{total:,.0f} FCFA**")
                if archived:
                    st.warning("**Employés archivés avec des consommations** — vérifier si dettes/retenues encore actives :")
                    for msg, total in sorted(archived.items(), key=lambda x: -x[1]):
                        st.markdown(f"- {msg} · **{total:,.0f} FCFA**")

        # ── Alertes Plafond Individuel ─────────────────────────────────────
        if rows_ytd and params.get("rates"):
            # Calcul YTD part_emp par employé
            ytd_emp_conso = defaultdict(lambda: {"emp": 0.0, "plafond": None, "client": ""})
            for r in rows_ytd:
                e = ytd_emp_conso[r["employee_name"]]
                e["emp"]     += r["part_emp"]
                e["client"]   = r["client"]
                if r.get("plafond_emp"):
                    e["plafond"] = r["plafond_emp"]

            emp_alerts = []
            for emp_name_a, ev in ytd_emp_conso.items():
                if ev["plafond"]:
                    pct = ev["emp"] / ev["plafond"] * 100
                    if pct >= 80:
                        emp_alerts.append((emp_name_a, ev["client"],
                                           ev["plafond"], ev["emp"], pct))

            if emp_alerts:
                with st.expander(
                    f"👤 {len(emp_alerts)} employé(s) proche(s) de leur plafond individuel",
                    expanded=len([a for a in emp_alerts if a[4] >= 100]) > 0
                ):
                    for emp_n, cli, plaf, consommé, pct in sorted(emp_alerts, key=lambda x: -x[4]):
                        restant = max(0, plaf - consommé)
                        if pct >= 100:
                            st.error(
                                f"🔴 **{emp_n}** ({cli}) — Plafond individuel **ATTEINT** : "
                                f"{consommé:,.0f} / {plaf:,.0f} FCFA ({pct:.1f} %)"
                            )
                        else:
                            st.warning(
                                f"🟡 **{emp_n}** ({cli}) — {pct:.1f} % consommé · "
                                f"Restant : {restant:,.0f} FCFA "
                                f"({consommé:,.0f} / {plaf:,.0f} FCFA)"
                            )

        # ── Alertes Provision Santé ────────────────────────────────────────
        if rows_ytd and params.get("rates"):
            provision_clients = {
                k: v for k, v in params["rates"].items()
                if v.get("modele") == "provision" and v.get("plafond")
            }
            if provision_clients:
                # Calcul YTD part soc par client
                ytd_soc_by_client = defaultdict(float)
                for r in rows_ytd:
                    ytd_soc_by_client[r["client"]] += r["part_soc"]

                alerts = []
                for client, rate_info in provision_clients.items():
                    plafond = rate_info["plafond"]
                    # Matching insensible à la casse
                    ytd_soc = 0
                    for k, v in ytd_soc_by_client.items():
                        if k.lower() == client.lower():
                            ytd_soc = v
                            break
                    pct = ytd_soc / plafond * 100 if plafond > 0 else 0
                    if pct >= 80:
                        alerts.append((client, plafond, ytd_soc, pct))

                if alerts:
                    with st.expander(f"🏥 {len(alerts)} alerte(s) Provision Santé", expanded=True):
                        for client, plafond, ytd_soc, pct in sorted(alerts, key=lambda x: -x[3]):
                            restant = max(0, plafond - ytd_soc)
                            if pct >= 100:
                                st.error(
                                    f"🔴 **{client}** — Plafond **DÉPASSÉ** : "
                                    f"{ytd_soc:,.0f} / {plafond:,.0f} FCFA ({pct:.1f} %)"
                                )
                            else:
                                st.warning(
                                    f"🟡 **{client}** — {pct:.1f} % consommé · "
                                    f"Restant : {restant:,.0f} FCFA "
                                    f"({ytd_soc:,.0f} / {plafond:,.0f} FCFA)"
                                )

        # ── Génération fichiers ───────────────────────────────────────────
        with st.spinner("Génération des fichiers…"):
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:

                # Export global (avec onglets YTD si disponibles)
                global_xls = build_global_excel(rows, period_label,
                                                 rows_ytd=rows_ytd, year=year)
                zf.writestr(f"Export_Global_{period_label.replace(' ', '_')}.xlsx",
                            global_xls)

                # ── Rapport Provision Santé ───────────────────────────────
                if rows_ytd:
                    prov_xls = build_provision_excel(rows_ytd, params, period_label, year)
                    zf.writestr(
                        f"Provision_Sante_YTD_{year}.xlsx",
                        prov_xls
                    )

                # ── Rapports clients (refacturation) ──────────────────────
                # 1. Récapitulatif multi-onglets (synthèse + un onglet/client)
                recap_xls = build_clients_recap_excel(rows, period_label)
                zf.writestr(
                    f"Clients/Recap_Clients_{period_label.replace(' ', '_')}.xlsx",
                    recap_xls
                )
                # 2. Un fichier Excel par client
                by_client_rows = defaultdict(list)
                for r in rows:
                    by_client_rows[r["client"]].append(r)
                for client_name, c_rows in by_client_rows.items():
                    safe_client = re.sub(r'[\\/*?:"<>|]', "_", client_name)
                    client_xls = build_client_excel(client_name, c_rows, period_label)
                    zf.writestr(
                        f"Clients/{safe_client}.xlsx",
                        client_xls
                    )

                # Fichiers individuels + bons de consommation
                by_emp = defaultdict(list)
                for r in rows:
                    by_emp[r["employee_name"]].append(r)

                for emp_name, emp_rows in by_emp.items():
                    safe_name = re.sub(r'[\\/*?:"<>|]', "_", emp_name)
                    emp_client = emp_rows[0]["client"] if emp_rows else ""

                    # Excel individuel
                    ind_xls = build_individual_excel(emp_name, emp_rows, period_label)
                    zf.writestr(f"Individuels/{safe_name}.xlsx", ind_xls)

                    # Relevé automatique .docx (sans template)
                    releve = build_releve_employe_docx(
                        emp_name, emp_client, emp_rows, period_label
                    )
                    zf.writestr(f"Relevés/{safe_name}.docx", releve)

                    # Bons de consommation (si templates fournis)
                    if tpl_emp:
                        tpl_emp.seek(0)
                        bon_emp = fill_bon_template(
                            tpl_emp.read(), emp_name, emp_rows,
                            period_label, "Employé"
                        )
                        zf.writestr(f"Bons_Employe/{safe_name}_bon_employe.docx", bon_emp)

                    if tpl_empeur:
                        tpl_empeur.seek(0)
                        bon_empeur = fill_bon_template(
                            tpl_empeur.read(), emp_name, emp_rows,
                            period_label, "Employeur"
                        )
                        zf.writestr(f"Bons_Employeur/{safe_name}_bon_employeur.docx",
                                    bon_empeur)

        zip_buf.seek(0)

        # ── Résumé mensuel ────────────────────────────────────────────────
        total_global = sum(r["montant_total"] for r in rows)
        total_soc    = sum(r["part_soc"]      for r in rows)
        total_emp    = sum(r["part_emp"]      for r in rows)

        st.subheader(f"📊 {period_label} — Résumé mensuel")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Employés concernés",  len(by_emp))
        c2.metric("Total consommations", f"{total_global:,.0f} FCFA")
        c3.metric("Part Société",        f"{total_soc:,.0f} FCFA")
        c4.metric("Part Employé(e)",     f"{total_emp:,.0f} FCFA")

        # ── Résumé YTD ────────────────────────────────────────────────────
        if rows_ytd and month > 1:
            ytd_global = sum(r["montant_total"] for r in rows_ytd)
            ytd_soc    = sum(r["part_soc"]      for r in rows_ytd)
            ytd_emp    = sum(r["part_emp"]      for r in rows_ytd)

            st.subheader(f"📅 {ytd_label} — Cumul YTD")
            d1, d2, d3, d4 = st.columns(4)
            ytd_emps = len({r["employee_name"] for r in rows_ytd})
            d1.metric("Employés (YTD)",    ytd_emps)
            d2.metric("Total YTD",         f"{ytd_global:,.0f} FCFA",
                      delta=f"+{ytd_global - total_global:,.0f} vs mois")
            d3.metric("Part Soc. YTD",     f"{ytd_soc:,.0f} FCFA")
            d4.metric("Part Emp. YTD",     f"{ytd_emp:,.0f} FCFA")

        st.success(f"✅ {len(by_emp)} fichiers individuels générés")

        # ── Téléchargement ────────────────────────────────────────────────
        st.download_button(
            label="📥 Télécharger tous les rapports (.zip)",
            data=zip_buf,
            file_name=f"AutoSante_{period_label.replace(' ', '_')}.zip",
            mime="application/zip",
            use_container_width=True,
            type="primary",
        )

    st.divider()

    # ── Phase 2 : Retenues sur salaire ────────────────────────────────────
    st.divider()
    st.subheader("💰 Phase 2 — Retenues sur salaire")
    st.caption(
        "Calcule, pour chaque employé(e), la retenue mensuelle à opérer sur salaire "
        "(Principe B : 15 % du total dû, min 10 000 FCFA si dette < 40 000 FCFA) "
        "et met à jour les soldes cumulatifs."
    )

    with st.expander("ℹ️ Comment ça marche ?", expanded=False):
        st.markdown(
            """
**Formule de calcul (Principe B — validé mars 2026) :**

| Variable | Description |
|---|---|
| `dette M-1` | Solde reporté du mois précédent (feuille *Retenues* du Google Sheet) |
| `conso M` | Part employé(e) calculée ci-dessus pour le mois sélectionné |
| `total dû` | `dette M-1 + conso M` |
| `retenue M` | 15 % × total dû — min 10 000 FCFA si total < 40 000 FCFA |
| `solde fin de mois` | `total dû − retenue M` |

**Règle par client** : une colonne *Règle_Retenue* (col 11) dans "Taux Clients" permet de paramétrer une règle différente par client :
`15%` · `total` · `1/3` · `1/4` · `1/5` · `fixe:25000`

**Google Sheets → feuille "Retenues"** (même classeur que Taux Clients) :

| Employé Nom | Client | Solde M-1 (FCFA) |
|---|---|---|
| DUPONT Marie | ERoCo | 55 250 |
| … | … | … |

Après génération, l'onglet **MAJ Soldes** du fichier Excel contient les nouveaux soldes à coller dans cette feuille.
            """
        )

    col_r1, col_r2 = st.columns([2, 1])
    with col_r1:
        st.info(
            "Les retenues sont calculées à partir des données du mois **déjà généré** ci-dessus. "
            "Sélectionnez le même mois et cliquez sur Générer les rapports d'abord."
        )
    with col_r2:
        gen_retenues = st.button(
            "💰 Générer le plan de retenues",
            type="primary",
            use_container_width=True,
            disabled="rows" not in st.session_state,
            help="Nécessite d'avoir généré les rapports du mois d'abord."
        )

    if gen_retenues or "retenue_rows" in st.session_state:
        # ── 1. Fetch salary attachments depuis Odoo ───────────────────────
        with st.spinner("Lecture des retenues actives dans Odoo…"):
            attachments_odoo = fetch_salary_attachments(uid, models)

        n_att = len(attachments_odoo)
        if n_att == 0:
            st.warning(
                "Aucun `hr.salary.attachment` de type 'Retraits Santé' trouvé dans Odoo "
                f"pour la société #{COMPANY_ID}. Vérifiez le modèle ou le filtre."
            )
        else:
            st.success(f"✅ {n_att} retenues actives chargées depuis Odoo")

        # ── 2. Aperçu rapide des attachments Odoo ────────────────────────
        with st.expander(f"📂 {n_att} retenues actives dans Odoo (avant calcul)", expanded=False):
            import pandas as pd
            if attachments_odoo:
                df_att = pd.DataFrame([
                    {
                        "Employé":           a["employee_name"],
                        "Retenue mensuelle": a["monthly_amount"],
                        "Total dû (Odoo)":   a["total_amount"],
                        "Déjà prélevé":      a["paid_amount"],
                        "Restant":           a["remaining"],
                        "Depuis":            a["date_start"],
                    }
                    for a in attachments_odoo.values()
                ]).sort_values("Restant", ascending=False)
                st.dataframe(
                    df_att.style.format({
                        "Retenue mensuelle": "{:,.0f}",
                        "Total dû (Odoo)":   "{:,.0f}",
                        "Déjà prélevé":      "{:,.0f}",
                        "Restant":           "{:,.0f}",
                    }).apply(lambda col: [
                        f"background-color: {'#FF6B35' if v > 100000 else '#FFB085' if v > 40000 else '#FFE0CC' if v > 0 else '#E2EFDA'}"
                        for v in col
                    ], subset=["Restant"]),
                    use_container_width=True, hide_index=True
                )

        # ── 3. Calcul des retenues ────────────────────────────────────────
        rows_for_retenues = st.session_state.get("rows", [])
        if not rows_for_retenues:
            st.error("Aucune donnée de conso disponible. Générez d'abord les rapports du mois.")
        else:
            year_r  = st.session_state.get("year",  year)
            month_r = st.session_state.get("month", month)
            with st.spinner("Calcul et réconciliation…"):
                ret_rows = compute_retenues(rows_for_retenues, attachments_odoo,
                                            st.session_state.get("params", {}),
                                            year=year_r, month=month_r)
                st.session_state["retenue_rows"] = ret_rows

            n_maj    = sum(1 for r in ret_rows if r["statut_attachment"] == "maj")
            n_new    = sum(1 for r in ret_rows if r["statut_attachment"] == "nouveau")
            n_soldé  = sum(1 for r in ret_rows if r["solde_fin_mois"] == 0)
            t_retenu = sum(r["new_retenue"]    for r in ret_rows)
            t_solde  = sum(r["solde_fin_mois"] for r in ret_rows)
            t_conso  = sum(r["conso_mois"]     for r in ret_rows)

            rc1, rc2, rc3, rc4, rc5 = st.columns(5)
            rc1.metric("✏️ Mises à jour",     n_maj)
            rc2.metric("🆕 À créer dans Odoo", n_new)
            rc3.metric("Total à retenir M",   f"{t_retenu:,.0f} FCFA")
            rc4.metric("Dette portée M+1",    f"{t_solde:,.0f} FCFA")
            rc5.metric("Conso M (part emp.)", f"{t_conso:,.0f} FCFA")

            # ── Tableau analyse ───────────────────────────────────────────
            with st.expander("📋 Détail complet", expanded=True):
                df_ret = pd.DataFrame([{
                    "Statut":            r["statut_attachment"],
                    "Employé":           r["employee_name"],
                    "Client":            r["client"],
                    "Retenue actuelle":  r["old_monthly"],
                    "Encours M-1":       r["dette_mn1"],
                    "Nlle Conso M":      r["conso_mois"],
                    "Nouveau total":     r["new_total"],
                    "Règle":             r["regle"],
                    "Nouvelle retenue":  r["new_retenue"],
                    "Solde fin de mois": r["solde_fin_mois"],
                } for r in ret_rows])
                st.dataframe(
                    df_ret.style.format({
                        "Retenue actuelle":  "{:,.0f}",
                        "Encours M-1":       "{:,.0f}",
                        "Nlle Conso M":      "{:,.0f}",
                        "Nouveau total":     "{:,.0f}",
                        "Nouvelle retenue":  "{:,.0f}",
                        "Solde fin de mois": "{:,.0f}",
                    }).apply(lambda col: [
                        f"background-color: {'#FFDAD9' if v > 100000 else '#FFE0CC' if v > 40000 else '#FFF9C4' if v > 0 else '#E2EFDA'}"
                        for v in col
                    ], subset=["Solde fin de mois"]),
                    use_container_width=True, hide_index=True
                )

            # ── Téléchargement ────────────────────────────────────────────
            period_label_r = st.session_state.get("period_label", period_label)
            # year_r / month_r déjà définis plus haut
            ret_xls = build_retenues_excel(ret_rows, period_label_r, year_r, month_r)

            # Avertissement anti-doublon AVANT le bouton de téléchargement
            if n_new > 0:
                st.warning(
                    f"⚠️ **{n_new} nouvelle(s) retenue(s) à créer** (onglet 'Import Odoo NEW').  \n"
                    "Cet onglet est à importer **une seule fois** dans Odoo.  \n"
                    "Après import, **rechargez la page et régénérez** le plan pour confirmer "
                    f"que ces {n_new} employé(s) passent bien en '✏️ MAJ' au prochain calcul.  \n"
                    "Si un employé est dans **plusieurs lots de paie** ce mois, la retenue "
                    "sera déduite à chaque bulletin — pensez à ajuster le champ "
                    "'Nombre d'échéances' dans Odoo si nécessaire."
                )

            col_dl1, col_dl2 = st.columns([3, 1])
            with col_dl1:
                st.download_button(
                    label=f"📥 Télécharger Retenues_{period_label_r.replace(' ', '_')}.xlsx",
                    data=ret_xls,
                    file_name=f"Retenues_{period_label_r.replace(' ', '_')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary",
                    use_container_width=True,
                )
            with col_dl2:
                if st.button("🔄 Forcer rechargement Odoo", use_container_width=True,
                             help="Vide le cache (5 min) et relit hr.salary.attachment depuis Odoo — "
                                  "utile après un import pour confirmer la prise en compte."):
                    fetch_salary_attachments.clear()
                    st.rerun()

            st.caption(
                "**Import Odoo MAJ** → met à jour `monthly_amount` et `total_amount` "
                "des retenues existantes (idempotent via `.id`).  \n"
                "**Import Odoo NEW** → crée de nouveaux `hr.salary.attachment` "
                "(description = 'Retrait Santé MM/AAAA' pour détecter les doublons)."
            )

        # ── Note Studio Odoo ──────────────────────────────────────────────
        with st.expander("🔧 Champs Pôle Santé sur la fiche employé (Studio)", expanded=False):
            st.markdown("""
**Pour afficher l'encours et la dernière retenue directement sur la fiche employé Odoo :**

1. Ouvrir **Studio** → modèle **Employé** (`hr.employee`) → onglet **Pôle Santé**
2. Ajouter un champ **Calculé** (type Monétaire) nommé `x_studio_encours_sante` :
   ```python
   # Logique (champ compute Python dans Studio)
   for rec in self:
       atts = self.env['hr.salary.attachment'].search([
           ('employee_id', '=', rec.id),
           ('state', 'in', ['open', 'running']),
           ('other_input_type_id.name', 'ilike', 'Santé'),
       ])
       rec.x_studio_encours_sante = sum(
           (a.total_amount - a.paid_amount) for a in atts
       )
   ```
   Label : **Encours total dettes santé** — Mode : **Lecture seule**

3. Ajouter un champ **Calculé** (type Monétaire) nommé `x_studio_derniere_retenue` :
   ```python
   for rec in self:
       att = self.env['hr.salary.attachment'].search([
           ('employee_id', '=', rec.id),
           ('state', 'in', ['open', 'running']),
           ('other_input_type_id.name', 'ilike', 'Santé'),
       ], order='date_start desc', limit=1)
       rec.x_studio_derniere_retenue = att.monthly_amount if att else 0
   ```
   Label : **Dernière retenue sur paie** — Mode : **Lecture seule**

> 💡 Ces champs seront visibles par les RH dans la fiche de chaque employé pour aider à expliquer les déductions sur fiche de paie.
            """)



# ── PHASE 2 : RETENUES SUR SALAIRE ───────────────────────────────────────

def fetch_retenues_sheet(sheet_id: str) -> dict:  # kept for compatibility — no longer used
    """
    Lit la feuille 'Retenues' du Google Sheet partagé.
    Structure attendue :
      Col 0 = Employé Nom  |  Col 1 = Client  |  Col 2 = Solde M-1 (FCFA)

    Retourne : {nom_employé: solde_mn1_float}
    """
    url = (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}"
        f"/gviz/tq?tqx=out:csv&sheet={urllib.parse.quote('Retenues')}"
    )
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "AutoSante/2.2"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read().decode("utf-8")
        rows = list(csv.reader(raw.splitlines()))
    except Exception as e:
        st.warning(f"⚠️ Impossible de lire la feuille 'Retenues' : {e} — soldes M-1 à zéro.")
        return {}

    result = {}
    for row in rows[1:]:  # skip header
        if not row or not row[0].strip():
            continue
        nom = row[0].strip()
        try:
            solde = float(str(row[2]).replace(",", "").replace(" ", "").strip()) if len(row) > 2 and row[2].strip() else 0.0
        except (ValueError, TypeError):
            solde = 0.0
        result[nom] = solde
    return result


def _calc_retenue(total_du: float, regle: str = "15%") -> float:
    """
    Applique la règle de retenue sur le total dû (dette M-1 + conso M).

    Règles acceptées :
      "15%"         → Principe B : 15% du total, min 10 000 FCFA si total < 40 000
      "total"       → Retrait en une fois
      "1/3"         → Un tiers
      "1/4"         → Un quart
      "1/5"         → Un cinquième
      "fixe:25000"  → Montant fixe mensuel (ex: 25 000 FCFA)
    """
    if total_du <= 0:
        return 0.0
    r = regle.strip().lower()
    if r == "total":
        return round(total_du, 0)
    elif r == "1/3":
        return round(total_du / 3, 0)
    elif r == "1/4":
        return round(total_du / 4, 0)
    elif r == "1/5":
        return round(total_du / 5, 0)
    elif r.startswith("fixe:"):
        try:
            return min(round(float(r.split(":")[1]), 0), total_du)
        except (ValueError, IndexError):
            pass
    # Par défaut : Principe B — 15%, min 10 000 FCFA si dette < 40 000
    retenue = total_du * 0.15
    if total_du < 40_000:
        retenue = max(retenue, 10_000.0)
    return min(round(retenue, 0), total_du)


def compute_retenues(rows_month: list, attachments_odoo: dict, params: dict,
                     year: int = None, month: int = None) -> list:
    """
    Réconcilie la consommation du mois (part_emp par employé) avec les
    hr.salary.attachment actifs dans Odoo.

    Trois cas :
      A) Attachment existant + conso M     → mise à jour monthly + total
      B) Attachment existant + conso M = 0 → portée de dette, recalcul retenue
      C) Conso M mais pas d'attachment     → nouvelle ligne à créer dans Odoo

    Retourne une liste de dicts, triée par client puis nom.
    """
    rates = params.get("rates", {})
    # Label période pour la description des nouveaux attachments
    _period_tag = f"{month:02d}/{year}" if year and month else ""

    # Agréger part_emp par employé depuis les lignes du mois
    by_emp: dict[str, dict] = {}
    for r in rows_month:
        name   = r["employee_name"]
        client = r["client"]
        if name not in by_emp:
            by_emp[name] = {"client": client, "conso_mois": 0.0}
        by_emp[name]["conso_mois"] += r["part_emp"]

    # Ajouter les employés avec attachment mais aucune conso ce mois
    for emp_name, att in attachments_odoo.items():
        if emp_name not in by_emp:
            by_emp[emp_name] = {"client": "", "conso_mois": 0.0}

    result = []
    for emp_name, info in by_emp.items():
        client     = info["client"]
        conso_mois = info["conso_mois"]
        att        = attachments_odoo.get(emp_name)

        # Dette M-1 = solde restant dans Odoo (total_amount - paid_amount)
        if att:
            dette_mn1       = att["remaining"]
            old_monthly     = att["monthly_amount"]
            odoo_id         = att["odoo_id"]
            description     = att["description"]
            date_start      = att["date_start"]
            input_type      = att["input_type"]
            statut_attachment = "maj"    # mise à jour
        else:
            dette_mn1       = 0.0
            old_monthly     = 0.0
            odoo_id         = None
            # La description inclut la période pour détecter les doublons visuellement
            description     = f"Retrait Santé {_period_tag}".strip() if _period_tag else "Retrait Santé"
            date_start      = ""
            input_type      = "Retraits Santé"
            statut_attachment = "nouveau"   # à créer dans Odoo

        new_total = dette_mn1 + conso_mois   # nouveau total_amount à pousser dans Odoo

        # Règle de retenue par client
        rate  = _find_rate(client, rates) if client else None
        regle = rate.get("regle_retenue", "15%") if rate else "15%"

        new_retenue   = _calc_retenue(new_total, regle)
        solde_fin     = max(0.0, new_total - new_retenue)

        result.append({
            # ── Identifiants Odoo ───────────────────────────────────────
            "odoo_id":          odoo_id,
            "description":      description,
            "date_start":       date_start,
            "input_type":       input_type,
            "statut_attachment": statut_attachment,   # "maj" | "nouveau"
            # ── Données employé ─────────────────────────────────────────
            "employee_name":    emp_name,
            "client":           client,
            # ── Calcul retenue ───────────────────────────────────────────
            "conso_mois":       round(conso_mois,   0),
            "dette_mn1":        round(dette_mn1,    0),   # remaining Odoo
            "new_total":        round(new_total,    0),   # → nouveau total_amount Odoo
            "old_monthly":      round(old_monthly,  0),   # retenue avant
            "regle":            regle,
            "new_retenue":      round(new_retenue,  0),   # → nouveau monthly_amount Odoo
            "solde_fin_mois":   round(solde_fin,    0),   # restant après retenue
        })

    return sorted(result, key=lambda x: (x["statut_attachment"], x["client"], x["employee_name"]))


def build_retenues_excel(retenue_rows: list, period_label: str,
                         year: int, month: int) -> bytes:
    """
    Génère Retenues_{MOIS}.xlsx avec 3 onglets :
      1. "Analyse {mois}"    : tableau complet — style Import Mars
      2. "Import Odoo MAJ"   : CSV prêt à importer dans Odoo (.id, monthly_amount, total_amount)
      3. "Import Odoo NEW"   : lignes à créer (employés sans attachment existant)
    """
    import datetime as _dt
    wb  = openpyxl.Workbook()

    fill_new  = PatternFill("solid", fgColor="DAEEF3")   # bleu pâle = nouveau
    fill_zero = PatternFill("solid", fgColor="E2EFDA")   # vert = soldé
    fill_red  = PatternFill("solid", fgColor="FFDAD9")   # rouge = dette élevée

    # ── Onglet 1 : Analyse ────────────────────────────────────────────────
    ws = wb.active
    ws.title = f"Analyse {period_label}"

    ws.merge_cells("A1:K1")
    ws["A1"] = f"Plan de retenues sur salaire — {period_label}  (source : Odoo hr.salary.attachment)"
    ws["A1"].font      = Font(bold=True, size=13, color="FFFFFF", name="Calibri")
    ws["A1"].fill      = PatternFill("solid", fgColor="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[1].height = 24

    hdrs = ["Statut", "Employé(e)", "Client",
            "Retenue actuelle\n(Odoo)", "Encours M-1\n(Odoo)",
            "Nlle Conso M\n(part emp.)",
            "Nouveau total\ndû", "Règle",
            "Nouvelle retenue\nM", "Solde fin\nde mois",
            "Δ Retenue"]
    ws.append(hdrs)
    _style_header(ws, 2, len(hdrs))
    ws.row_dimensions[2].height = 36

    t_old_m = t_dette = t_conso = t_new_total = t_new_ret = t_solde = 0.0

    for i, r in enumerate(retenue_rows):
        is_new = r["statut_attachment"] == "nouveau"
        delta  = r["new_retenue"] - r["old_monthly"]
        ws.append([
            "🆕 À créer" if is_new else "✏️ MAJ",
            r["employee_name"], r["client"],
            r["old_monthly"] if not is_new else "",
            r["dette_mn1"],  r["conso_mois"],
            r["new_total"],  r["regle"],
            r["new_retenue"], r["solde_fin_mois"],
            delta if not is_new else "",
        ])
        dr = i + 3
        rf = fill_new if is_new else (fill_zero if r["solde_fin_mois"] == 0 else fill_red)
        for ci in range(1, len(hdrs) + 1):
            cell = ws.cell(row=dr, column=ci)
            if rf:
                cell.fill = rf
            cell.font      = Font(name="Calibri", size=10)
            cell.alignment = Alignment(horizontal="left")
        for ci in [4, 5, 6, 7, 9, 10, 11]:
            ws.cell(row=dr, column=ci).number_format = '#,##0" FCFA"'

        t_old_m   += r["old_monthly"]
        t_dette   += r["dette_mn1"]
        t_conso   += r["conso_mois"]
        t_new_total += r["new_total"]
        t_new_ret += r["new_retenue"]
        t_solde   += r["solde_fin_mois"]

    lr = 3 + len(retenue_rows)
    ws.append(["", f"TOTAL ({len(retenue_rows)} employés)", "",
               t_old_m, t_dette, t_conso, t_new_total, "",
               t_new_ret, t_solde, t_new_ret - t_old_m])
    _style_header(ws, lr, len(hdrs), bg="2E75B6")
    for ci in [4, 5, 6, 7, 9, 10, 11]:
        ws.cell(row=lr, column=ci).number_format = '#,##0" FCFA"'

    for col, w in zip(range(1, len(hdrs) + 1),
                      [12, 34, 22, 18, 16, 16, 16, 12, 18, 18, 14]):
        ws.column_dimensions[get_column_letter(col)].width = w
    ws.freeze_panes = "B3"

    # ── Onglet 2 : Import Odoo MAJ (lignes existantes à mettre à jour) ────
    ws2 = wb.create_sheet("Import Odoo MAJ")
    ws2.merge_cells("A1:E1")
    ws2["A1"] = (f"Import Odoo — mise à jour retenues — {period_label}  "
                 f"(coller dans Odoo > Salaires > Acomptes & Saisies)")
    ws2["A1"].font      = Font(bold=True, size=11, color="FFFFFF", name="Calibri")
    ws2["A1"].fill      = PatternFill("solid", fgColor="375623")
    ws2["A1"].alignment = Alignment(horizontal="center")

    # En-tête compatible import Odoo (utilise .id = database integer ID)
    hdrs2 = [".id", "description", "employee_ids",
             "monthly_amount", "total_amount"]
    ws2.append(hdrs2)
    _style_header(ws2, 2, len(hdrs2), bg="375623")

    maj_rows = [r for r in retenue_rows if r["statut_attachment"] == "maj"]
    for i, r in enumerate(maj_rows):
        ws2.append([
            r["odoo_id"],
            r["description"],
            r["employee_name"],
            int(r["new_retenue"]),
            int(r["new_total"]),
        ])
        _style_data(ws2, i + 3, len(hdrs2), alt=(i % 2 == 1))
        for ci in [4, 5]:
            ws2.cell(row=i + 3, column=ci).number_format = '#,##0'

    for col, w in zip(range(1, len(hdrs2) + 1), [14, 28, 38, 18, 18]):
        ws2.column_dimensions[get_column_letter(col)].width = w
    ws2.freeze_panes = "A3"

    # ── Onglet 3 : Import Odoo NEW (nouvelles lignes à créer) ─────────────
    ws3 = wb.create_sheet("Import Odoo NEW")
    ws3.merge_cells("A1:F1")
    ws3["A1"] = f"Import Odoo — nouvelles retenues à créer — {period_label}"
    ws3["A1"].font      = Font(bold=True, size=11, color="FFFFFF", name="Calibri")
    ws3["A1"].fill      = PatternFill("solid", fgColor="1F4E79")
    ws3["A1"].alignment = Alignment(horizontal="center")

    # Ligne d'avertissement anti-doublon
    ws3.merge_cells("A2:F2")
    ws3["A2"] = (
        "⚠️  IMPORTER UNE SEULE FOIS par employé par période. "
        "Après import, relancer l'app — ces employés doivent passer en onglet 'Import Odoo MAJ'. "
        "Si ce n'est pas le cas, vérifier que le type de saisie 'Retraits Santé' existe dans Odoo."
    )
    ws3["A2"].font      = Font(bold=True, size=10, color="7B2C00", name="Calibri")
    ws3["A2"].fill      = PatternFill("solid", fgColor="FFE0CC")
    ws3["A2"].alignment = Alignment(horizontal="left", wrap_text=True)
    ws3.row_dimensions[2].height = 32

    hdrs3 = ["description", "employee_ids", "monthly_amount",
             "total_amount", "date_start", "other_input_type_id"]
    ws3.append(hdrs3)
    _style_header(ws3, 3, len(hdrs3))

    new_rows = [r for r in retenue_rows if r["statut_attachment"] == "nouveau"]
    # date_start = 1er jour du mois traité (year/month passés via build_retenues_excel)
    first_of_month = _dt.date(year, month, 1).isoformat() if year and month else _dt.date.today().replace(day=1).isoformat()
    for i, r in enumerate(new_rows):
        ws3.append([
            r["description"],
            r["employee_name"],
            int(r["new_retenue"]),
            int(r["new_total"]),
            first_of_month,
            "Retraits Santé",
        ])
        dr3 = i + 4   # +4 car : titre(1) + avertissement(2) + en-têtes(3)
        for ci in range(1, len(hdrs3) + 1):
            cell = ws3.cell(row=dr3, column=ci)
            cell.fill      = fill_new
            cell.font      = Font(name="Calibri", size=10)
            cell.alignment = Alignment(horizontal="left")
        for ci in [3, 4]:
            ws3.cell(row=dr3, column=ci).number_format = '#,##0'

    if not new_rows:
        ws3.append(["(aucune nouvelle retenue ce mois)"])

    for col, w in zip(range(1, len(hdrs3) + 1), [28, 38, 18, 18, 14, 20]):
        ws3.column_dimensions[get_column_letter(col)].width = w
    ws3.freeze_panes = "A4"

    # Couleur onglets
    for sheet in wb.sheetnames:
        wb[sheet].sheet_properties.tabColor = "1F4E79"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    main()
